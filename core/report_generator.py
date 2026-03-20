"""Report generation for normalized ThermoAnalyzer result records."""

from __future__ import annotations

import csv
import io
import datetime
import os
import re
from typing import Any, Mapping, Optional, Union

from core.batch_runner import normalize_batch_summary_rows, summarize_batch_outcomes
from core.chemical_formula_formatting import format_chemical_formula_text
from core.mechanism_rules import tga_mechanism_signals
from core.processing_schema import ensure_processing_payload
from core.xrd_reference_dossier import XRD_NO_VISUAL_ASSET_NOTE
from core.result_serialization import flatten_result_records, partition_results_by_status, split_valid_results
from core.scientific_sections import (
    build_tga_scientific_narrative,
    condense_warning_limitations,
    normalize_report_text,
    normalize_scientific_context,
    scientific_context_to_report_sections,
)
from core.literature_partitioning import partition_reference_ids
from core.xrd_display import xrd_candidate_display_name
from utils.reference_data import find_nearest_reference

try:
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError as exc:  # pragma: no cover
    raise ImportError(
        "python-docx is required for DOCX report generation. Install it with: pip install python-docx"
    ) from exc


_MAIN_METADATA_KEYS = (
    "sample_name",
    "sample_mass",
    "heating_rate",
    "instrument",
    "vendor",
    "atmosphere",
    "display_name",
    "file_name",
    "import_confidence",
    "import_confidence_label",
    "inferred_analysis_type",
)

_APPENDIX_METADATA_KEYWORDS = (
    "hash",
    "delimiter",
    "decimal",
    "header",
    "row",
    "start",
    "import",
    "warning",
    "heuristic",
    "inferred",
    "parser",
    "dialect",
    "encoding",
)

_PAPER_DISPLAY_LABEL_OVERRIDES = {
    "tga_polymers_comparison.xlsx": "PMMA sample",
    "tga_cuso4_5h2o_dehydration.csv": "CuSO4·5H2O sample",
    "tga_caco3_decomposition.csv": "CaCO3 sample",
}

_BULLET_SECTION_TITLES = {
    "Scientific Interpretation",
    "Primary Scientific Interpretation",
    "Evidence Supporting This Interpretation",
    "Alternative Explanations",
    "Uncertainty and Methodological Limits",
    "Recommended Follow-Up Experiments",
    "Literature Comparison",
    "Relevant References",
    "Alternative or Non-Validating References",
    "Supporting References",
    "Contradictory or Alternative References",
    "Recommended Follow-Up Literature Checks",
}
_MARKDOWN_LINK_PATTERN = re.compile(r"\[([^\]]+)\]\((https?://[^)\s]+)\)")
_BARE_URL_PATTERN = re.compile(r"(https?://[^\s)]+)")


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    for child in tr_pr.findall(qn("w:tblHeader")):
        tr_pr.remove(child)
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _docx_section_text_width_inches(section) -> float:
    width = float(section.page_width.inches) - float(section.left_margin.inches) - float(section.right_margin.inches)
    return max(width, 1.0)


def _estimate_docx_column_widths(headers: list[str], rows: list[list[Any]], *, total_width_inches: float) -> list[float]:
    if not headers:
        return []
    max_lengths: list[int] = []
    for column_index, header in enumerate(headers):
        max_len = len(str(header or ""))
        for row in rows[:40]:
            if column_index >= len(row):
                continue
            max_len = max(max_len, len(_format_value(row[column_index])))
        max_lengths.append(max_len)
    weights = [max(1.0, min(6.0, length / 12.0)) for length in max_lengths]
    weight_total = sum(weights) or float(len(weights))
    return [total_width_inches * weight / weight_total for weight in weights]


def _apply_docx_table_widths(table, widths: list[float]) -> None:
    if not widths:
        return
    table.autofit = False
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            if index >= len(widths):
                continue
            cell.width = Inches(widths[index])


def _set_docx_section_orientation(section, *, landscape: bool) -> None:
    if landscape:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        if section.page_width < section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width
    else:
        section.orientation = WD_ORIENTATION.PORTRAIT
        if section.page_width > section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(normalize_report_text(text), level=level)


def _presentation_text(value: Any) -> str:
    return format_chemical_formula_text(normalize_report_text(value))


def _append_docx_hyperlink(paragraph, *, url: str, text: str) -> None:
    clean_url = normalize_report_text(url)
    clean_text = _presentation_text(text)
    if not clean_url or not clean_text:
        return
    relationship_id = paragraph.part.relate_to(clean_url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    run_properties.append(style)
    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = clean_text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _append_docx_text_with_links(paragraph, text: str) -> None:
    clean_text = _presentation_text(text)
    if not clean_text:
        return

    def append_plain(segment: str) -> None:
        cursor = 0
        for url_match in _BARE_URL_PATTERN.finditer(segment):
            if url_match.start() > cursor:
                paragraph.add_run(segment[cursor:url_match.start()])
            url = url_match.group(1)
            _append_docx_hyperlink(paragraph, url=url, text=url)
            cursor = url_match.end()
        if cursor < len(segment):
            paragraph.add_run(segment[cursor:])

    cursor = 0
    for match in _MARKDOWN_LINK_PATTERN.finditer(clean_text):
        if match.start() > cursor:
            append_plain(clean_text[cursor:match.start()])
            _append_docx_hyperlink(paragraph, url=match.group(2), text=match.group(1))
        cursor = match.end()
    if cursor < len(clean_text):
        append_plain(clean_text[cursor:])


def _doi_url(value: Any) -> str:
    cleaned = normalize_report_text(value)
    if not cleaned:
        return ""
    if cleaned.lower().startswith(("http://", "https://")):
        return cleaned
    return f"https://doi.org/{cleaned}"


def _citation_link_markup(citation: Mapping[str, Any]) -> str:
    doi = normalize_report_text(citation.get("doi") or "")
    url = normalize_report_text(citation.get("url") or "")
    if doi:
        return normalize_report_text(f"DOI: [{doi}]({_doi_url(doi)}).")
    if url:
        return normalize_report_text(f"Link: [Open paper]({url}).")
    return ""


def _add_key_value_table(doc: Document, data: dict) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = normalize_report_text("Parameter")
    hdr_cells[1].text = normalize_report_text("Value")
    _set_repeat_table_header(table.rows[0])
    for cell in hdr_cells:
        _set_cell_bg(cell, "4472C4")
        for para in cell.paragraphs:
            run = para.runs[0] if para.runs else para.add_run(cell.text)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for key, value in data.items():
        row_cells = table.add_row().cells
        row_cells[0].text = normalize_report_text(key)
        row_cells[1].text = normalize_report_text(value)
    text_width = _docx_section_text_width_inches(doc.sections[-1])
    _apply_docx_table_widths(table, [text_width * 0.34, text_width * 0.66])


def _add_results_table(doc: Document, headers: list[str], rows: list[list]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = normalize_report_text(header)
        _set_cell_bg(hdr_cells[i], "4472C4")
        for para in hdr_cells[i].paragraphs:
            run = para.runs[0] if para.runs else para.add_run(hdr_cells[i].text)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _set_repeat_table_header(table.rows[0])

    for row_idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = normalize_report_text(value)
            if row_idx % 2 == 1:
                _set_cell_bg(row_cells[i], "DCE6F1")
    _apply_docx_table_widths(
        table,
        _estimate_docx_column_widths(
            headers,
            rows,
            total_width_inches=_docx_section_text_width_inches(doc.sections[-1]),
        ),
    )


def _format_value(value):
    if value is None:
        return "N/A"
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, list):
        return "; ".join(str(item) for item in value) if value else "N/A"
    if isinstance(value, dict):
        return ", ".join(f"{key}={_format_value(item)}" for key, item in value.items()) if value else "N/A"
    try:
        return f"{float(value):.4f}"
    except (TypeError, ValueError):
        return normalize_report_text(value)


def _truncate_report_text(value: Any, *, max_chars: int = 320) -> str:
    text = normalize_report_text(value)
    if len(text) <= max_chars:
        return text
    return f"{text[: max_chars - 1]}…"


def _compact_collection_value(value: Any, *, max_items: int = 6, max_chars: int = 320) -> str:
    if value in (None, "", [], {}):
        return "N/A"
    if isinstance(value, Mapping):
        items = list(value.items())
        fragments: list[str] = []
        for key, item in items[:max_items]:
            if isinstance(item, Mapping):
                token = f"{key}=[{len(item)} fields]"
            elif isinstance(item, list):
                token = f"{key}=[{len(item)} items]"
            else:
                token = f"{key}={_format_value(item)}"
            fragments.append(_truncate_report_text(token, max_chars=max_chars))
        if len(items) > max_items:
            fragments.append(f"…(+{len(items) - max_items})")
        return _truncate_report_text("; ".join(fragments), max_chars=max_chars)
    if isinstance(value, list):
        if not value:
            return "N/A"
        if all(not isinstance(item, (list, dict, tuple, set)) for item in value):
            preview = "; ".join(_format_value(item) for item in value[:max_items])
            if len(value) > max_items:
                preview = f"{preview}; …(+{len(value) - max_items})"
            return _truncate_report_text(preview, max_chars=max_chars)
        return _truncate_report_text(f"[{len(value)} items]", max_chars=max_chars)
    return _truncate_report_text(_format_value(value), max_chars=max_chars)


def _compact_xrd_evidence_for_table(value: Any, *, max_chars: int = 320) -> str:
    if not isinstance(value, Mapping):
        return _compact_collection_value(value, max_chars=max_chars)
    tokens: list[str] = []
    for key in (
        "shared_peak_count",
        "weighted_overlap_score",
        "coverage_ratio",
        "mean_delta_position",
        "unmatched_major_peak_count",
    ):
        if value.get(key) not in (None, ""):
            tokens.append(f"{key}={_format_value(value.get(key))}")
    if value.get("matched_peak_pairs") is not None:
        tokens.append(f"matched_peak_pairs={len(value.get('matched_peak_pairs') or [])}")
    if value.get("unmatched_observed_peaks") is not None:
        tokens.append(f"unmatched_observed={len(value.get('unmatched_observed_peaks') or [])}")
    if value.get("unmatched_reference_peaks") is not None:
        tokens.append(f"unmatched_reference={len(value.get('unmatched_reference_peaks') or [])}")
    if not tokens:
        return _compact_collection_value(value, max_chars=max_chars)
    return _truncate_report_text("; ".join(tokens), max_chars=max_chars)


def _format_table_cell_value(*, analysis_type: str, header: str, value: Any, max_chars: int = 320) -> str:
    header_key = str(header or "").strip().lower()
    normalized_type = str(analysis_type or "").upper()
    if normalized_type == "XRD" and header_key == "evidence":
        return _compact_xrd_evidence_for_table(value, max_chars=max_chars)
    return _compact_collection_value(value, max_chars=max_chars)


def _record_title(record: dict) -> str:
    dataset_key = record.get("dataset_key")
    if dataset_key:
        return normalize_report_text(f"{record['analysis_type']} - {dataset_key}")
    return normalize_report_text(record["analysis_type"])


def _record_headers(record: dict) -> list[str]:
    rows = record.get("rows") or []
    if not rows:
        return []
    first_row = rows[0]
    return list(first_row.keys())


def _humanize_key(key: str) -> str:
    replacements = {"id": "ID", "utc": "UTC", "dsc": "DSC", "tga": "TGA", "dtg": "DTG"}
    words = str(key).replace("_", " ").split()
    return " ".join(replacements.get(word.lower(), word.capitalize()) for word in words)


def _table_payload(payload: dict | None) -> dict[str, str]:
    table_rows: dict[str, str] = {}
    for key, value in (payload or {}).items():
        if value in (None, "", [], {}):
            continue
        table_rows[_humanize_key(key)] = _format_value(value)
    return table_rows


def _compact_table_payload(payload: Mapping[str, Any] | None, *, max_chars: int = 220) -> dict[str, str]:
    table_rows: dict[str, str] = {}
    for key, value in (payload or {}).items():
        if value in (None, "", [], {}):
            continue
        table_rows[_humanize_key(key)] = _compact_collection_value(value, max_chars=max_chars)
    return table_rows


def _safe_float(value: Any) -> float | None:
    try:
        if value in (None, ""):
            return None
        return float(value)
    except (TypeError, ValueError):
        return None


def _format_number(value: Any, *, digits: int = 2) -> str:
    numeric = _safe_float(value)
    if numeric is None:
        return "N/A"
    return f"{numeric:.{digits}f}"


def _dataset_label(dataset_key: str | None, datasets: dict) -> str:
    if not dataset_key:
        return "N/A"
    dataset = datasets.get(dataset_key)
    if dataset is None:
        return dataset_key
    metadata = getattr(dataset, "metadata", {}) or {}
    return normalize_report_text(metadata.get("display_name") or metadata.get("sample_name") or metadata.get("file_name") or dataset_key)


def _paper_display_label(dataset_key: str | None, datasets: dict, *, record: dict | None = None) -> str:
    def normalized_lookup_key(value: Any) -> str:
        raw = str(value or "").strip()
        if not raw:
            return ""
        return os.path.basename(raw).lower()

    def paperize_fallback(value: Any) -> str:
        raw = str(value or "").strip()
        if not raw:
            return ""
        base = os.path.basename(raw)
        stem, ext = os.path.splitext(base)
        if ext.lower() in {".csv", ".xlsx", ".xls", ".txt"}:
            cleaned = stem.replace("_", " ").replace("-", " ")
            cleaned = " ".join(cleaned.split())
            lowered = cleaned.lower()
            for prefix in ("tga ", "dsc ", "dta ", "kissinger ", "ofw ", "friedman "):
                if lowered.startswith(prefix):
                    cleaned = cleaned[len(prefix):]
                    break
            if cleaned:
                return normalize_report_text(f"{cleaned} sample")
        return normalize_report_text(raw)

    if dataset_key and dataset_key in datasets:
        metadata = getattr(datasets[dataset_key], "metadata", {}) or {}
        summary = (record or {}).get("summary") or {}
        candidates = [
            metadata.get("display_name"),
            metadata.get("sample_name"),
            summary.get("sample_name"),
            metadata.get("file_name"),
            dataset_key,
        ]
    else:
        metadata = (record or {}).get("metadata") or {}
        summary = (record or {}).get("summary") or {}
        candidates = [
            metadata.get("display_name"),
            metadata.get("sample_name"),
            summary.get("sample_name"),
            metadata.get("file_name"),
            dataset_key,
            (record or {}).get("dataset_key"),
        ]

    for value in candidates:
        if value in (None, ""):
            continue
        lookup = normalized_lookup_key(value)
        if lookup in _PAPER_DISPLAY_LABEL_OVERRIDES:
            return normalize_report_text(_PAPER_DISPLAY_LABEL_OVERRIDES[lookup])
    for value in candidates:
        if value not in (None, ""):
            return paperize_fallback(value)
    return "Unnamed dataset"


def _xrd_best_candidate_name(summary: dict[str, Any]) -> str:
    return str(
        summary.get("top_candidate_display_name_unicode")
        or summary.get("top_phase_display_name_unicode")
        or xrd_candidate_display_name(summary, target="unicode")
        or "N/A"
    )


def _xrd_top_phase_name(summary: dict[str, Any]) -> str:
    return str(
        summary.get("top_phase_display_name_unicode")
        or summary.get("top_candidate_display_name_unicode")
        or xrd_candidate_display_name(summary, target="unicode")
        or "N/A"
    )


def _xrd_best_candidate_score(summary: dict[str, Any]) -> Any:
    value = summary.get("top_candidate_score")
    if value in (None, ""):
        value = summary.get("top_phase_score")
    return value


def _xrd_has_best_candidate(summary: dict[str, Any]) -> bool:
    return _xrd_best_candidate_name(summary) != "N/A"


def _main_conditions_payload(dataset_key: str, dataset) -> dict[str, str]:
    metadata = getattr(dataset, "metadata", {}) or {}
    def value_or_not_recorded(value: Any) -> str:
        if value in (None, "", [], {}):
            return "Not recorded"
        return _format_value(value)

    return {
        "Dataset": value_or_not_recorded(_dataset_label(dataset_key, {dataset_key: dataset})),
        "Source File": value_or_not_recorded(metadata.get("file_name") or metadata.get("display_name")),
        "Sample Name": value_or_not_recorded(metadata.get("sample_name")),
        "Sample Mass": value_or_not_recorded(metadata.get("sample_mass")),
        "Heating Rate": value_or_not_recorded(metadata.get("heating_rate")),
        "Instrument": value_or_not_recorded(metadata.get("instrument")),
        "Vendor": value_or_not_recorded(metadata.get("vendor")),
        "Atmosphere": value_or_not_recorded(metadata.get("atmosphere")),
    }


def _appendix_dataset_metadata(metadata: dict | None) -> dict[str, str]:
    metadata = metadata or {}
    appendix_payload = {}
    for key, value in metadata.items():
        if value in (None, "", [], {}):
            continue
        lower = str(key).lower()
        if key in _MAIN_METADATA_KEYS:
            continue
        if any(token in lower for token in _APPENDIX_METADATA_KEYWORDS):
            appendix_payload[key] = value
    return _table_payload(appendix_payload)


def _record_key_results(record: dict) -> dict[str, str]:
    summary = record.get("summary") or {}
    analysis_type = str(record.get("analysis_type") or "").upper()
    if analysis_type == "TGA":
        keep = ("step_count", "total_mass_loss_percent", "residue_percent", "sample_name", "sample_mass", "heating_rate")
    elif analysis_type == "DSC":
        keep = (
            "peak_count",
            "glass_transition_count",
            "tg_midpoint",
            "tg_onset",
            "tg_endset",
            "delta_cp",
            "sample_name",
            "sample_mass",
            "heating_rate",
        )
    elif analysis_type in {"FTIR", "RAMAN"}:
        keep = (
            "match_status",
            "top_match_id",
            "top_match_name",
            "top_match_score",
            "confidence_band",
            "candidate_count",
            "library_provider",
            "library_package",
            "library_version",
            "library_sync_mode",
            "library_cache_status",
            "library_access_mode",
            "library_request_id",
            "library_result_source",
            "library_provider_scope",
            "library_offline_limited_mode",
            "caution_code",
            "caution_message",
            "sample_name",
            "sample_mass",
            "heating_rate",
        )
    elif analysis_type == "XRD":
        return _table_payload(
            {
                "Accepted Match Status": summary.get("match_status"),
                "Best Candidate": _xrd_best_candidate_name(summary),
                "Best Candidate Score": _xrd_best_candidate_score(summary),
                "Confidence Band": summary.get("confidence_band"),
                "Shared Peaks": summary.get("top_candidate_shared_peak_count"),
                "Coverage Ratio": summary.get("top_candidate_coverage_ratio"),
                "Caution Code": summary.get("caution_code"),
                "Best Candidate Reason Below Threshold": summary.get("top_candidate_reason_below_threshold"),
            }
        )
    else:
        keep = tuple(summary.keys())
    return _table_payload({key: summary.get(key) for key in keep})


def _record_metric_snapshot(record: dict) -> str:
    summary = record.get("summary") or {}
    analysis_type = str(record.get("analysis_type") or "").upper()
    if analysis_type == "TGA":
        return ", ".join(
            [
                f"total mass loss {_format_number(summary.get('total_mass_loss_percent'))}%",
                f"residue {_format_number(summary.get('residue_percent'))}%",
                f"step count {_format_value(summary.get('step_count'))}",
            ]
        )
    if analysis_type == "DSC":
        snapshot = [f"peak count {_format_value(summary.get('peak_count'))}"]
        if summary.get("tg_midpoint") is not None:
            snapshot.append(f"Tg midpoint {_format_number(summary.get('tg_midpoint'))} °C")
        return ", ".join(snapshot)
    if analysis_type in {"FTIR", "RAMAN"}:
        match_status = str(summary.get("match_status") or "not_recorded")
        confidence_band = str(summary.get("confidence_band") or "not_recorded")
        top_score = _format_number(summary.get("top_match_score"), digits=3)
        top_match = summary.get("top_match_name") or summary.get("top_match_id") or "N/A"
        library_label = summary.get("library_package") or summary.get("library_provider") or "embedded"
        result_source = summary.get("library_result_source") or "unknown_source"
        if match_status.lower() == "matched":
            return ", ".join(
                [
                    f"top match {top_match}",
                    f"score {top_score}",
                    f"confidence {confidence_band}",
                    f"library {library_label}",
                    f"source {result_source}",
                ]
            )
        caution = summary.get("caution_code") or "spectral_no_match"
        return ", ".join(
            [
                f"match status {match_status}",
                f"top score {top_score}",
                f"library {library_label}",
                f"caution {caution}",
                f"source {result_source}",
            ]
        )
    if analysis_type == "XRD":
        match_status = str(summary.get("match_status") or "not_recorded")
        confidence_band = str(summary.get("confidence_band") or "not_recorded")
        top_score = _format_number(_xrd_best_candidate_score(summary), digits=3)
        top_phase = _xrd_top_phase_name(summary)
        best_candidate = _xrd_best_candidate_name(summary)
        library_label = summary.get("library_package") or summary.get("library_provider") or "embedded"
        result_source = summary.get("library_result_source") or "unknown_source"
        if match_status.lower() == "matched":
            return ", ".join(
                [
                    f"top phase {top_phase}",
                    f"score {top_score}",
                    f"confidence {confidence_band}",
                    f"library {library_label}",
                    f"source {result_source}",
                ]
            )
        if _xrd_has_best_candidate(summary):
            caution = summary.get("caution_code") or "xrd_no_match"
            shared_peaks = _format_value(summary.get("top_candidate_shared_peak_count"))
            reason = summary.get("top_candidate_reason_below_threshold") or "evidence below acceptance threshold"
            return ", ".join(
                [
                    f"best candidate {best_candidate}",
                    f"score {top_score}",
                    f"shared peaks {shared_peaks}",
                    f"accepted match status {match_status}",
                    f"caution {caution}",
                    f"reason {reason}",
                    f"source {result_source}",
                ]
            )
        caution = summary.get("caution_code") or "xrd_no_match"
        return ", ".join(
            [
                f"match status {match_status}",
                f"top phase score {top_score}",
                f"library {library_label}",
                f"caution {caution}",
                f"source {result_source}",
            ]
        )

    parts = []
    for key, value in summary.items():
        if key in {"sample_name", "sample_mass", "heating_rate"}:
            continue
        parts.append(f"{_humanize_key(key)} {_format_value(value)}")
        if len(parts) >= 3:
            break
    return ", ".join(parts) if parts else "Key metrics not available"


def _record_confidence_note(record: dict) -> str:
    grouped = condense_warning_limitations(
        record.get("scientific_context"),
        validation=record.get("validation"),
        max_data_warnings=2,
        max_method_limits=2,
    )
    data_warnings = grouped.get("data_completeness_warnings") or []
    method_limits = grouped.get("methodological_limitations") or []
    if data_warnings:
        return data_warnings[0]
    if method_limits:
        return method_limits[0]

    status = ((record.get("validation") or {}).get("status") or "").lower()
    if status == "pass":
        return "Validation checks did not flag material data-quality issues."
    if status == "warn":
        return "Validation generated warnings; interpret with moderate caution."
    if status == "fail":
        return "Validation failed; interpretation is not suitable for definitive conclusions."
    return "No explicit confidence constraints were recorded."


def _record_compact_rows(record: dict, *, max_rows: int = 5) -> tuple[list[str], list[list[str]]] | None:
    rows = record.get("rows") or []
    if not rows:
        return None
    analysis_type = str(record.get("analysis_type") or "").upper()
    if analysis_type == "TGA":
        return None
    if analysis_type == "XRD":
        return _xrd_candidate_mini_table(record, max_rows=max_rows)

    headers = _record_headers(record)
    if not headers:
        return None

    if len(rows) <= max_rows:
        return (
            headers,
            [[_format_value(row.get(header)) for header in headers] for row in rows],
        )

    top_headers = headers[: min(5, len(headers))]
    top_rows = [[_format_value(row.get(header)) for header in top_headers] for row in rows[:max_rows]]
    return top_headers, top_rows


def _xrd_candidate_mini_table(record: dict, *, max_rows: int = 5) -> tuple[list[str], list[list[str]]] | None:
    if str(record.get("analysis_type") or "").upper() != "XRD":
        return None
    rows = record.get("rows") or []
    if not rows:
        return None

    summary = record.get("summary") or {}
    matrix: list[list[str]] = []
    for row in rows[:max_rows]:
        evidence = dict(row.get("evidence") or {})
        matrix.append(
            [
                _format_value(row.get("rank")),
                xrd_candidate_display_name(row, target="unicode") or _format_value(row.get("candidate_name")),
                _format_number(row.get("normalized_score"), digits=3),
                _format_value(row.get("confidence_band") or summary.get("match_status")),
                _format_value(evidence.get("shared_peak_count")),
                _format_number(evidence.get("coverage_ratio"), digits=3),
            ]
        )
    if not matrix:
        return None
    return ["Rank", "Candidate", "Score", "Status", "Shared Peaks", "Coverage"], matrix


def _tga_major_events(record: dict, *, limit: int = 3) -> list[list[str]]:
    if str(record.get("analysis_type") or "").upper() != "TGA":
        return []

    scored: list[tuple[float, dict]] = []
    for row in (record.get("rows") or []):
        if not isinstance(row, dict):
            continue
        loss = _safe_float(row.get("mass_loss_percent"))
        scored.append((loss if loss is not None else float("-inf"), row))

    scored.sort(key=lambda item: item[0], reverse=True)
    top = [row for _, row in scored[:limit] if row]

    output = []
    for idx, row in enumerate(top, start=1):
        output.append(
            [
                f"Event {idx}",
                _format_number(row.get("midpoint_temperature")),
                _format_number(row.get("mass_loss_percent")),
                _format_number(row.get("residual_percent")),
            ]
        )
    return output


def _record_full_rows(record: dict) -> tuple[list[str], list[list[str]]] | None:
    headers = _record_headers(record)
    if not headers:
        return None
    analysis_type = str(record.get("analysis_type") or "").upper()
    rows = [
        [
            _format_table_cell_value(
                analysis_type=analysis_type,
                header=header,
                value=row.get(header),
            )
            for header in headers
        ]
        for row in (record.get("rows") or [])
    ]
    if not rows:
        return None
    return headers, rows


def _select_record_for_dataset(records: list[dict], dataset_key: str, analysis_type: str | None) -> dict | None:
    candidates = [record for record in records if record.get("dataset_key") == dataset_key]
    if not candidates:
        return None
    if analysis_type:
        normalized = analysis_type.upper()
        filtered = [record for record in candidates if str(record.get("analysis_type") or "").upper() == normalized]
        if filtered:
            stable = [record for record in filtered if record.get("status") == "stable"]
            return stable[0] if stable else filtered[0]
    stable = [record for record in candidates if record.get("status") == "stable"]
    return stable[0] if stable else candidates[0]


def _comparison_dataset_label(dataset_key: str, datasets: dict, records: list[dict], analysis_type: str | None) -> str:
    record = _select_record_for_dataset(records, dataset_key, analysis_type)
    return _paper_display_label(dataset_key, datasets, record=record)


def _comparison_missing_metadata(selected: list[str], datasets: dict, *, analysis_type: str | None = None) -> list[str]:
    normalized = str(analysis_type or "").upper()
    if normalized == "TGA":
        required = (("heating_rate", "heating rate"), ("atmosphere", "atmosphere"))
    elif normalized == "XRD":
        required = (("xrd_wavelength_angstrom", "wavelength"), ("xrd_axis_unit", "xrd axis unit"))
    else:
        required = (("heating_rate", "heating rate"),)
    missing: list[str] = []
    for key, label in required:
        if any(not ((getattr(datasets.get(dataset_key), "metadata", {}) or {}).get(key)) for dataset_key in selected):
            missing.append(label)
    return missing


def _comparison_limitation_sentence(*, missing_metadata: list[str], excluded_labels: list[str], partial: bool) -> str:
    fragments = []
    if excluded_labels:
        fragments.append(
            "comparable metrics are not yet available for "
            + (excluded_labels[0] if len(excluded_labels) == 1 else f"{', '.join(excluded_labels[:-1])}, and {excluded_labels[-1]}")
        )
    if missing_metadata:
        fragments.append(
            "key metadata are incomplete, including "
            + (missing_metadata[0] if len(missing_metadata) == 1 else f"{', '.join(missing_metadata[:-1])}, and {missing_metadata[-1]}")
        )
    if not fragments:
        return ""
    scope = "partial rather than comprehensive" if partial else "comparative rather than definitive"
    return f"Because {'; '.join(fragments)}, this interpretation should be treated as {scope}."


def _build_tga_comparison_interpretation(
    metrics: list[dict[str, Any]],
    *,
    excluded_labels: list[str],
    missing_metadata: list[str],
) -> str:
    def behavior_from_metric(metric: dict[str, Any]) -> str:
        signals = metric.get("signals") or {}
        class_inference = signals.get("material_class_inference") or {}
        class_type = str(class_inference.get("material_class") or "")
        mass_balance = signals.get("mass_balance_assessment") or {}
        if str(mass_balance.get("status") or "") in {"strong_match", "plausible_match"}:
            pathway = str(mass_balance.get("pathway") or "").strip()
            if pathway:
                return f"mass balance is consistent with near-complete {pathway}"
            if class_type == "hydrate_salt":
                return "mass balance is consistent with near-complete dehydration to an expected anhydrous residue"
            if class_type == "carbonate_inorganic":
                return "mass balance is consistent with near-complete decarbonation to a stable oxide-rich residue"
            if class_type == "hydroxide_to_oxide":
                return "mass balance is consistent with dehydroxylation toward a stable oxide residue"
            if class_type == "oxalate_multistage_inorganic":
                return "mass balance is consistent with a multistage gas-loss conversion to a stable residue"
            if class_type == "generic_inorganic_salt_or_mineral":
                return "mass balance is consistent with conversion to an expected stable inorganic residue"
        if float(metric["mass_loss"]) >= 90.0 and float(metric["residue"]) <= 10.0:
            return "metrics indicate extensive decomposition over the recorded range"
        return "metrics indicate partial conversion with retained residue, requiring class-specific context for mechanistic assignment"

    def pathway_from_metric(metric: dict[str, Any]) -> str | None:
        signals = metric.get("signals") or {}
        class_inference = signals.get("material_class_inference") or {}
        class_type = str(class_inference.get("material_class") or "")
        mass_balance = signals.get("mass_balance_assessment") or {}
        if str(mass_balance.get("status") or "") not in {"strong_match", "plausible_match"}:
            return None
        specific_pathway = str(mass_balance.get("pathway") or "").strip()
        if specific_pathway:
            return specific_pathway
        if class_type == "hydrate_salt":
            return "dehydration toward an anhydrous inorganic end-product"
        if class_type == "carbonate_inorganic":
            return "decarbonation toward an oxide-rich inorganic end-product"
        if class_type == "hydroxide_to_oxide":
            return "dehydroxylation toward an oxide end-product"
        if class_type == "oxalate_multistage_inorganic":
            return "multistage oxalate gas-loss conversion toward a stable inorganic end-product"
        if class_type == "generic_inorganic_salt_or_mineral":
            return "class-consistent conversion toward a stable inorganic end-product"
        return "class-consistent conversion toward a stable end-product"

    if len(metrics) >= 2:
        mass_losses = [float(item["mass_loss"]) for item in metrics]
        residues = [float(item["residue"]) for item in metrics]
        step_counts = [int(item["step_count"]) for item in metrics]
        highest_loss = max(metrics, key=lambda item: float(item["mass_loss"]))
        lowest_loss = min(metrics, key=lambda item: float(item["mass_loss"]))

        text = (
            f"Across the reportable TGA datasets, total mass loss spans {min(mass_losses):.2f}% to {max(mass_losses):.2f}%, "
            f"final residue spans {min(residues):.2f}% to {max(residues):.2f}%, and resolved step counts range from {min(step_counts)} to {max(step_counts)}. "
            f"{highest_loss['dataset']} shows the highest overall mass-loss extent, whereas {lowest_loss['dataset']} retains comparatively higher residue."
        )
        chemistry_entries = []
        for metric in metrics:
            pathway = pathway_from_metric(metric)
            if pathway:
                chemistry_entries.append((metric["dataset"], pathway))
        if len(chemistry_entries) >= 2:
            first_name, first_pathway = chemistry_entries[0]
            second_name, second_pathway = chemistry_entries[1]
            text = (
                f"{text} Chemistry-aware interpretation indicates {first_name} is consistent with {first_pathway}, "
                f"while {second_name} is consistent with {second_pathway}. "
                "The compared TGA datasets differ not only in total mass loss and final residue, but also in the chemistry "
                "of the expected final solid products. A higher residue may reflect retention of a stable inorganic end-product "
                "rather than incomplete conversion. Accordingly, these runs should be interpreted as different transformation "
                "pathways, not simply as different degrees of decomposition."
            )
        elif len(chemistry_entries) == 1:
            only_name, only_pathway = chemistry_entries[0]
            text = (
                f"{text} For {only_name}, class-aware mass-balance evidence is consistent with {only_pathway}. "
                "Even when residue differs across runs, higher residue does not automatically indicate less complete conversion; "
                "it can reflect retention of a different stable solid end-product."
            )
        else:
            text = (
                f"{text} Residue differences should be interpreted with pathway context: higher residue does not automatically "
                "mean less complete conversion and may reflect different stable end-products."
            )
    elif len(metrics) == 1:
        only = metrics[0]
        behavior = behavior_from_metric(only)
        pathway = pathway_from_metric(only)
        chemistry_tail = ""
        if pathway:
            chemistry_tail = (
                f" The inferred pathway is {pathway}; retained residue should therefore be interpreted with end-product chemistry "
                "rather than as a simple proxy for decomposition completeness."
            )
        text = (
            f"Within the current comparison workspace, only {only['dataset']} produced reportable TGA summary metrics. "
            f"That dataset shows total mass loss of {float(only['mass_loss']):.2f}%, final residue of {float(only['residue']):.2f}%, "
            f"and {int(only['step_count'])} resolved decomposition steps; {behavior}.{chemistry_tail}"
        )
    else:
        text = "None of the selected datasets currently provide reportable TGA comparison metrics."

    limitation = _comparison_limitation_sentence(
        missing_metadata=missing_metadata,
        excluded_labels=excluded_labels,
        partial=len(metrics) <= 1 or bool(excluded_labels),
    )
    return f"{text} {limitation}".strip()


def _build_generic_comparison_interpretation(
    analysis_type: str,
    metrics: list[dict[str, Any]],
    *,
    excluded_labels: list[str],
    missing_metadata: list[str],
) -> str:
    if len(metrics) >= 2:
        text = (
            f"The {analysis_type} comparison shows measurable differences across the selected datasets. "
            "Review key metrics together with method context before drawing definitive mechanistic conclusions."
        )
    elif len(metrics) == 1:
        text = f"Only one {analysis_type} dataset currently has reportable metrics, so cross-dataset interpretation remains partial."
    else:
        text = f"No {analysis_type} metrics were available for comparison interpretation."

    limitation = _comparison_limitation_sentence(
        missing_metadata=missing_metadata,
        excluded_labels=excluded_labels,
        partial=len(metrics) <= 1 or bool(excluded_labels),
    )
    return f"{text} {limitation}".strip()


def _build_comparison_payload(comparison_workspace: dict | None, datasets: dict, records: list[dict]) -> dict[str, Any] | None:
    comparison_workspace = comparison_workspace or {}
    selected = comparison_workspace.get("selected_datasets") or []
    if not selected:
        return None

    analysis_type = str(comparison_workspace.get("analysis_type") or "N/A")
    normalized_analysis = analysis_type.upper()

    overview = {
        "Analysis Type": analysis_type,
        "Compared Datasets": ", ".join(
            _comparison_dataset_label(dataset_key, datasets, records, normalized_analysis)
            for dataset_key in selected
        ),
        "Saved Figure": comparison_workspace.get("figure_key") or "Not recorded",
    }

    metric_headers: list[str]
    metric_rows: list[list[str]] = []
    reportable_metric_records: list[dict[str, Any]] = []
    excluded_dataset_labels: list[str] = []
    missing_metadata = _comparison_missing_metadata(selected, datasets, analysis_type=normalized_analysis)

    if normalized_analysis == "TGA":
        metric_headers = ["Dataset", "Total Mass Loss (%)", "Final Residue (%)", "Step Count"]
        for dataset_key in selected:
            dataset_name = _comparison_dataset_label(dataset_key, datasets, records, normalized_analysis)
            record = _select_record_for_dataset(records, dataset_key, normalized_analysis)
            summary = (record or {}).get("summary") or {}
            mass_loss = _safe_float(summary.get("total_mass_loss_percent"))
            residue = _safe_float(summary.get("residue_percent"))
            step_count_raw = summary.get("step_count")
            step_count = int(step_count_raw) if isinstance(step_count_raw, (int, float)) else None
            if mass_loss is None or residue is None or step_count is None:
                excluded_dataset_labels.append(dataset_name)
                continue
            payload = {
                "dataset": dataset_name,
                "mass_loss": mass_loss,
                "residue": residue,
                "step_count": step_count,
                "signals": tga_mechanism_signals(
                    summary,
                    (record or {}).get("rows") or [],
                    metadata=(record or {}).get("metadata") or {},
                    dataset_key=str(dataset_key),
                ),
            }
            reportable_metric_records.append(payload)
            metric_rows.append(
                [
                    dataset_name,
                    _format_number(mass_loss),
                    _format_number(residue),
                    _format_value(step_count),
                ]
            )
        interpretation = _build_tga_comparison_interpretation(
            reportable_metric_records,
            excluded_labels=excluded_dataset_labels,
            missing_metadata=missing_metadata,
        )
    else:
        metric_headers = ["Dataset", "Primary Metrics"]
        for dataset_key in selected:
            dataset_name = _comparison_dataset_label(dataset_key, datasets, records, normalized_analysis)
            record = _select_record_for_dataset(records, dataset_key, normalized_analysis)
            if record is None:
                excluded_dataset_labels.append(dataset_name)
                continue
            metric = _record_metric_snapshot(record or {})
            reportable_metric_records.append({"dataset": dataset_name})
            metric_rows.append([dataset_name, metric])
        interpretation = _build_generic_comparison_interpretation(
            analysis_type,
            reportable_metric_records,
            excluded_labels=excluded_dataset_labels,
            missing_metadata=missing_metadata,
        )

    appendix_overview = _table_payload(
        {
            "saved_at": comparison_workspace.get("saved_at"),
            "notes": comparison_workspace.get("notes"),
            "batch_run_id": comparison_workspace.get("batch_run_id"),
            "batch_template_label": comparison_workspace.get("batch_template_label"),
            "batch_template_id": comparison_workspace.get("batch_template_id"),
            "batch_completed_at": comparison_workspace.get("batch_completed_at"),
        }
    )

    return {
        "overview": _table_payload(overview),
        "metric_headers": metric_headers,
        "metric_rows": metric_rows,
        "interpretation": interpretation,
        "reportable_count": len(reportable_metric_records),
        "missing_metadata": missing_metadata,
        "excluded_labels": excluded_dataset_labels,
        "excluded_note": (
            "Excluded from metric comparison (pending or not reportable): "
            + ", ".join(excluded_dataset_labels)
            if excluded_dataset_labels
            else ""
        ),
        "appendix_overview": appendix_overview,
        "appendix_batch_rows": _comparison_batch_rows(comparison_workspace),
    }


def _build_executive_summary_intro(records: list[dict], datasets: dict, comparison_payload: dict[str, Any] | None) -> str:
    if not records:
        return "This report currently has no analyzable records."

    reportable_tga = [
        record for record in records
        if str(record.get("analysis_type") or "").upper() == "TGA"
        and _safe_float((record.get("summary") or {}).get("total_mass_loss_percent")) is not None
        and _safe_float((record.get("summary") or {}).get("residue_percent")) is not None
    ]
    main_line = "This report summarizes the processed thermal-analysis outputs for the selected datasets."
    if reportable_tga:
        lead = max(reportable_tga, key=lambda record: float((record.get("summary") or {}).get("total_mass_loss_percent") or 0.0))
        lead_summary = lead.get("summary") or {}
        lead_name = _dataset_label(lead.get("dataset_key"), datasets) if lead.get("dataset_key") else (lead.get("id") or "the leading dataset")
        main_line = (
            "This report summarizes thermogravimetric and related thermal-analysis outputs for the selected datasets. "
            f"Among the processed runs, {lead_name} shows total mass loss of approximately {float(lead_summary.get('total_mass_loss_percent')):.1f}% "
            f"with final residue near {float(lead_summary.get('residue_percent')):.1f}%."
        )

    limitation = "Interpretation is strongest for reportable runs and should be qualified by dataset completeness."
    if comparison_payload:
        limitation = _comparison_limitation_sentence(
            missing_metadata=comparison_payload.get("missing_metadata") or [],
            excluded_labels=comparison_payload.get("excluded_labels") or [],
            partial=(comparison_payload.get("reportable_count") or 0) <= 1 or bool(comparison_payload.get("excluded_labels")),
        ) or limitation
    return normalize_report_text(f"{main_line} {limitation}")


def _build_executive_summary_rows(records: list[dict], datasets: dict, comparison_payload: dict[str, Any] | None) -> list[list[str]]:
    rows: list[list[str]] = []

    for record in records:
        dataset_name = _dataset_label(record.get("dataset_key"), datasets)
        analysis_type = str(record.get("analysis_type") or "Analysis")
        status = str(record.get("status") or "unknown")
        interpretation_sections = scientific_context_to_report_sections(record.get("scientific_context"))
        interpretation_line = "No interpretation statement recorded."
        preferred_titles = ("Primary Scientific Interpretation", "Scientific Interpretation")
        for title, payload in interpretation_sections:
            if title in preferred_titles and isinstance(payload, dict) and payload:
                interpretation_line = str(next(iter(payload.values())))
                break
        rows.append(
            [
                dataset_name,
                f"{analysis_type} ({status})",
                _record_metric_snapshot(record),
                interpretation_line,
                _record_confidence_note(record),
            ]
        )

    if comparison_payload:
        excluded = comparison_payload.get("excluded_labels") or []
        compared = comparison_payload.get("reportable_count") or 0
        rows.append(
            [
                "Comparison Set",
                "Cross-dataset comparison",
                f"Reportable datasets: {compared}; pending/not reportable: {len(excluded)}",
                comparison_payload.get("interpretation") or "Comparison interpretation unavailable",
                comparison_payload.get("excluded_note") or "Comparative interpretation depends on metadata completeness and method parity.",
            ]
        )

    return rows


def _build_final_conclusion_paragraph(records: list[dict], comparison_payload: dict[str, Any] | None) -> str:
    if not records:
        return "No validated analysis results were available to support a final scientific conclusion."

    def classify_tga_behavior(record: dict, mass_loss: float, residue: float, step_count: str) -> str:
        signals = tga_mechanism_signals(
            record.get("summary") or {},
            record.get("rows") or [],
            metadata=record.get("metadata") or {},
            dataset_key=str(record.get("dataset_key") or ""),
        )
        class_inference = signals.get("material_class_inference") or {}
        class_type = str(class_inference.get("material_class") or "")
        mass_balance = signals.get("mass_balance_assessment") or {}
        pathway = str(mass_balance.get("pathway") or "").strip()
        if str(mass_balance.get("status") or "") in {"strong_match", "plausible_match"}:
            if pathway:
                return f"behavior consistent with {pathway}"
            if class_type == "hydrate_salt":
                return "near-complete dehydration to an expected stable anhydrous residue"
            if class_type == "carbonate_inorganic":
                return "near-complete decarbonation to a stable oxide-rich residue"
            if class_type == "hydroxide_to_oxide":
                return "dehydroxylation toward an expected stable oxide residue"
            if class_type == "oxalate_multistage_inorganic":
                return "a multistage gas-loss conversion toward a stable final residue"
            if class_type == "generic_inorganic_salt_or_mineral":
                return "conversion to an expected stable inorganic residue"
        step_value = _safe_float(step_count)
        if mass_loss >= 90.0 and residue <= 10.0:
            return "near-complete decomposition with minimal residue"
        if residue >= 40.0 or mass_loss <= 60.0:
            return "partial decomposition with substantial residue-forming behavior"
        if step_value is not None and step_value >= 3:
            return "multi-step decomposition with notable residue retention"
        return "partial decomposition with moderate residue retention"

    tga_records = []
    for record in records:
        if str(record.get("analysis_type") or "").upper() != "TGA":
            continue
        summary = record.get("summary") or {}
        mass_loss = _safe_float(summary.get("total_mass_loss_percent"))
        residue = _safe_float(summary.get("residue_percent"))
        if mass_loss is None or residue is None:
            continue
        tga_records.append((record, mass_loss, residue, _format_value(summary.get("step_count"))))

    tga_conclusion = ""
    if len(tga_records) >= 2:
        low_residue = min(tga_records, key=lambda item: item[2])
        high_residue = max(tga_records, key=lambda item: item[2])
        low_name = _paper_display_label(low_residue[0].get("dataset_key"), {}, record=low_residue[0])
        high_name = _paper_display_label(high_residue[0].get("dataset_key"), {}, record=high_residue[0])
        low_behavior = classify_tga_behavior(low_residue[0], low_residue[1], low_residue[2], low_residue[3])
        high_behavior = classify_tga_behavior(high_residue[0], high_residue[1], high_residue[2], high_residue[3])
        tga_conclusion = (
            f"In TGA scope, {low_name} shows {low_behavior} "
            f"(mass loss {low_residue[1]:.1f}%, residue {low_residue[2]:.1f}%, step count {low_residue[3]}), "
            f"whereas {high_name} shows {high_behavior} "
            f"(mass loss {high_residue[1]:.1f}%, residue {high_residue[2]:.1f}%, step count {high_residue[3]})."
        )
    elif len(tga_records) == 1:
        only = tga_records[0]
        only_name = _paper_display_label(only[0].get("dataset_key"), {}, record=only[0])
        only_behavior = classify_tga_behavior(only[0], only[1], only[2], only[3])
        tga_conclusion = (
            f"In TGA scope, {only_name} shows {only_behavior} "
            f"(mass loss {only[1]:.1f}%, residue {only[2]:.1f}%, step count {only[3]})."
        )

    families: list[str] = []
    for record in records:
        analysis = str(record.get("analysis_type") or "").upper()
        if analysis in {"KISSINGER", "OZAWA-FLYNN-WALL", "FRIEDMAN"}:
            label = "kinetics"
        elif analysis == "PEAK DECONVOLUTION":
            label = "peak deconvolution"
        elif analysis in {"DSC", "TGA", "DTA"}:
            label = analysis
        else:
            label = analysis.lower() or "other analyses"
        if label not in families:
            families.append(label)

    if not families:
        sentence = "The report compiles the available thermal-analysis outputs, but no reportable analysis family was available for synthesis."
    elif len(families) == 1:
        sentence = f"The report provides a focused synthesis of the {families[0]} findings and their methodological constraints."
    else:
        sentence = (
            "The report integrates findings across "
            + ", ".join(families[:-1])
            + f", and {families[-1]}, providing a multi-technique interpretation rather than a single-modality conclusion."
        )
    if tga_conclusion:
        sentence = f"{sentence} {tga_conclusion}"

    limitation = "Metadata limitations reduce mechanistic certainty and the conclusions should be treated as condition-dependent."
    if comparison_payload:
        limitation = _comparison_limitation_sentence(
            missing_metadata=comparison_payload.get("missing_metadata") or [],
            excluded_labels=comparison_payload.get("excluded_labels") or [],
            partial=(comparison_payload.get("reportable_count") or 0) <= 1 or bool(comparison_payload.get("excluded_labels")),
        ) or limitation
    return normalize_report_text(f"{sentence} {limitation}")


def _analysis_family_label(analysis_type: str | None) -> str:
    normalized = str(analysis_type or "").upper()
    if normalized in {"KISSINGER", "OZAWA-FLYNN-WALL", "FRIEDMAN"}:
        return "Kinetics"
    if normalized == "PEAK DECONVOLUTION":
        return "Peak Deconvolution"
    return normalized or "Unknown"


def _build_pdf_abstract_layout(records: list[dict], datasets: dict) -> tuple[str, list[str]]:
    if not records:
        return ("No analyzable records were available for abstract-level synthesis.", [])

    family_counts: dict[str, int] = {}
    for record in records:
        family = _analysis_family_label(record.get("analysis_type"))
        family_counts[family] = family_counts.get(family, 0) + 1
    family_phrase = ", ".join(f"{family} (n={count})" for family, count in family_counts.items())
    abstract = (
        "This report presents a scientific synthesis of the processed thermal-analysis records. "
        f"The analyzed families were {family_phrase}. "
        "Interpretations are evidence-linked and uncertainty-qualified at the analysis level."
    )

    def clean_sentence(text: str) -> str:
        cleaned = normalize_report_text(text).strip()
        while ".." in cleaned:
            cleaned = cleaned.replace("..", ".")
        cleaned = cleaned.rstrip(" ;,:")
        if cleaned and cleaned[-1] not in ".!?":
            cleaned = f"{cleaned}."
        return cleaned

    bullets: list[str] = []
    for record in records:
        dataset_name = _paper_display_label(record.get("dataset_key"), datasets, record=record)
        analysis_label = _analysis_family_label(record.get("analysis_type")).lower()
        key_findings = _record_metric_snapshot(record) or "key findings were not fully reportable"
        primary_line = ""
        sections = scientific_context_to_report_sections(record.get("scientific_context"))
        for title, payload in sections:
            if title == "Primary Scientific Interpretation" and isinstance(payload, dict) and payload:
                primary_line = str(next(iter(payload.values()))).split(".")[0]
                break
        if primary_line:
            bullet = f"{dataset_name} ({analysis_label}): {key_findings}; {primary_line}"
        else:
            bullet = f"{dataset_name} ({analysis_label}): {key_findings}"
        bullets.append(clean_sentence(bullet))
    return normalize_report_text(abstract), bullets[:3]


def _build_experimental_prose_block(dataset_key: str, dataset, datasets: dict) -> str:
    metadata = getattr(dataset, "metadata", {}) or {}
    sample = _paper_display_label(dataset_key, datasets)
    source_file = metadata.get("file_name") or metadata.get("display_name") or "Not recorded"
    instrument = metadata.get("instrument") or "Not recorded"
    heating_rate = metadata.get("heating_rate")
    atmosphere = metadata.get("atmosphere") or "Not recorded"
    heating_rate_text = f"{_format_value(heating_rate)}" if heating_rate not in (None, "", [], {}) else "Not recorded"
    missing = []
    for key, label in (
        ("file_name", "source file"),
        ("instrument", "instrument"),
        ("heating_rate", "heating rate"),
        ("atmosphere", "atmosphere"),
    ):
        if metadata.get(key) in (None, "", [], {}):
            missing.append(label)
    completeness = "Metadata completeness is adequate for primary interpretation."
    if missing:
        completeness = (
            "Metadata completeness is limited; not recorded fields include "
            + (missing[0] if len(missing) == 1 else f"{', '.join(missing[:-1])}, and {missing[-1]}")
            + "."
        )
    return normalize_report_text(
        f"{sample} was analyzed from source file {source_file}. "
        f"Instrument: {instrument}. Heating rate: {heating_rate_text}. Atmosphere: {atmosphere}. "
        f"{completeness}"
    )


def _paper_record_heading(record: dict, datasets: dict) -> str:
    analysis = _analysis_family_label(record.get("analysis_type"))
    label = _paper_display_label(record.get("dataset_key"), datasets, record=record)
    return normalize_report_text(f"{analysis}: {label}")


def _is_comparison_figure_caption(value: str) -> bool:
    lowered = str(value or "").lower()
    return "comparison workspace" in lowered or lowered.startswith("comparison ")


def _normalized_caption(value: Any) -> str:
    lowered = str(value or "").lower()
    chars = [char if char.isalnum() else " " for char in lowered]
    return " ".join("".join(chars).split())


def _record_identity_markers(record: dict) -> list[str]:
    metadata = record.get("metadata") or {}
    summary = record.get("summary") or {}
    dataset_key = record.get("dataset_key")
    raw_candidates = [
        dataset_key,
        os.path.basename(str(dataset_key or "")),
        metadata.get("file_name"),
        os.path.basename(str(metadata.get("file_name") or "")),
        metadata.get("display_name"),
        metadata.get("sample_name"),
        summary.get("sample_name"),
    ]
    markers: list[str] = []
    for candidate in raw_candidates:
        normalized = _normalized_caption(candidate)
        if not normalized:
            continue
        if normalized not in markers:
            markers.append(normalized)
        parts = [part for part in normalized.split() if len(part) >= 4]
        if len(parts) >= 2:
            phrase = " ".join(parts)
            if phrase not in markers:
                markers.append(phrase)
    return markers


def _record_identity_tokens(record: dict) -> set[str]:
    metadata = record.get("metadata") or {}
    summary = record.get("summary") or {}
    raw_candidates = [
        record.get("dataset_key"),
        metadata.get("file_name"),
        metadata.get("display_name"),
        metadata.get("sample_name"),
        summary.get("sample_name"),
    ]
    tokens: set[str] = set()
    for candidate in raw_candidates:
        for token in re.findall(r"[a-z0-9]+(?:[_-][a-z0-9]+)+", str(candidate or "").lower()):
            tokens.add(token)
    return tokens


def _caption_matches_record(caption: str, record: dict) -> bool:
    normalized_caption = _normalized_caption(caption)
    if not normalized_caption:
        return False
    for marker in _record_identity_markers(record):
        if len(marker) < 4:
            continue
        if marker in normalized_caption:
            return True
    return False


def _caption_explicitly_conflicts_record(caption: str, record: dict) -> bool:
    caption_tokens = set(re.findall(r"[a-z0-9]+(?:[_-][a-z0-9]+)+", str(caption or "").lower()))
    if not caption_tokens:
        return False
    record_tokens = _record_identity_tokens(record)
    if not record_tokens:
        return False
    return caption_tokens.isdisjoint(record_tokens)


def _record_figure_keys(record: dict) -> list[str]:
    keys = (record.get("artifacts") or {}).get("figure_keys")
    if isinstance(keys, list):
        return [str(item) for item in keys if item not in (None, "")]
    return []


def _record_primary_figure_key(record: dict) -> str | None:
    value = (record.get("artifacts") or {}).get("report_figure_key")
    if value in (None, ""):
        return None
    return str(value)


def select_record_figures(record: dict, figures: dict | None, used: set[str]) -> list[tuple[str, bytes]]:
    figures = figures or {}
    explicit: list[tuple[str, bytes]] = []
    matched: list[tuple[str, bytes]] = []
    safe_generic: list[tuple[str, bytes]] = []

    primary_key = _record_primary_figure_key(record)
    if (
        primary_key
        and primary_key in figures
        and primary_key not in used
        and not _is_comparison_figure_caption(primary_key)
        and not _caption_explicitly_conflicts_record(primary_key, record)
    ):
        used.add(primary_key)
        return [(primary_key, figures[primary_key])]

    preferred_keys = _record_figure_keys(record)
    for key in preferred_keys:
        if _is_comparison_figure_caption(key):
            continue
        if key in figures and key not in used:
            if _caption_explicitly_conflicts_record(key, record):
                continue
            # Artifact-linked figures are record-level intent; prioritize them.
            explicit.append((key, figures[key]))

    if explicit:
        for caption, _ in explicit:
            used.add(caption)
        return explicit

    for caption, png_bytes in figures.items():
        if caption in used:
            continue
        if _is_comparison_figure_caption(caption):
            continue
        if _caption_matches_record(caption, record):
            matched.append((caption, png_bytes))
            continue
        if not _caption_explicitly_conflicts_record(caption, record):
            safe_generic.append((caption, png_bytes))
    if matched:
        for caption, _ in matched:
            used.add(caption)
        return matched

    if safe_generic:
        for caption, _ in safe_generic:
            used.add(caption)
        return safe_generic
    return matched


def _figures_for_record(record: dict, figures: dict | None, used: set[str]) -> list[tuple[str, bytes]]:
    """Backward-compatible alias for record-scoped figure routing."""
    return select_record_figures(record, figures, used)


def select_comparison_figures(
    figures: dict | None,
    used: set[str],
    *,
    comparison_workspace: dict | None = None,
) -> list[tuple[str, bytes]]:
    figures = figures or {}
    comparison_workspace = comparison_workspace or {}
    matched: list[tuple[str, bytes]] = []

    preferred_key = comparison_workspace.get("figure_key")
    if preferred_key and preferred_key in figures and preferred_key not in used and _is_comparison_figure_caption(str(preferred_key)):
        matched.append((str(preferred_key), figures[preferred_key]))
        used.add(str(preferred_key))

    for caption, png_bytes in figures.items():
        if caption in used:
            continue
        if not _is_comparison_figure_caption(caption):
            continue
        matched.append((caption, png_bytes))
        used.add(caption)
    return matched


def _record_main_mini_table(record: dict) -> tuple[list[str], list[list[str]]] | None:
    summary = record.get("summary") or {}
    analysis = str(record.get("analysis_type") or "").upper()
    if analysis == "TGA":
        payload = [
            ["Total Mass Loss (%)", _format_number(summary.get("total_mass_loss_percent"))],
            ["Final Residue (%)", _format_number(summary.get("residue_percent"))],
            ["Step Count", _format_value(summary.get("step_count"))],
        ]
    elif analysis == "DSC":
        payload = [
            ["Peak Count", _format_value(summary.get("peak_count"))],
            ["Tg Midpoint (°C)", _format_number(summary.get("tg_midpoint"))],
            ["Delta Cp", _format_number(summary.get("delta_cp"))],
        ]
    elif analysis == "XRD":
        payload = [
            ["Peak Count", _format_value(summary.get("peak_count"))],
            ["Accepted Match Status", _format_value(summary.get("match_status"))],
            ["Best Candidate", _xrd_best_candidate_name(summary)],
            ["Best Candidate Score", _format_number(_xrd_best_candidate_score(summary), digits=3)],
            ["Caution Code", _format_value(summary.get("caution_code"))],
        ]
    else:
        payload = []
        for key, value in summary.items():
            if key in {"sample_name", "sample_mass", "heating_rate"}:
                continue
            payload.append([_humanize_key(key), _format_value(value)])
            if len(payload) >= 3:
                break
    payload = [row for row in payload if row[1] not in {"N/A", "None"}]
    if not payload:
        return None
    return ["Outcome", "Value"], payload


def _processing_step(processing: dict | None, key: str) -> dict:
    processing = processing or {}
    nested = processing.get("signal_pipeline") or {}
    if key in nested and isinstance(nested[key], dict):
        return nested[key]

    nested = processing.get("analysis_steps") or {}
    if key in nested and isinstance(nested[key], dict):
        return nested[key]

    value = processing.get(key)
    return value if isinstance(value, dict) else {}


def _format_processing_step(payload: dict | None) -> str:
    payload = payload or {}
    if not payload:
        return "Not recorded"

    method = payload.get("method")
    details = []
    for key, value in payload.items():
        if key == "method" or value in (None, "", [], {}):
            continue
        details.append(f"{_humanize_key(key)}={_format_value(value)}")

    if method and details:
        return f"{method} ({'; '.join(details)})"
    if method:
        return str(method)
    if details:
        return "; ".join(details)
    return "Recorded"


def _reference_visibility(record: dict) -> str:
    analysis_type = record.get("analysis_type")
    if analysis_type not in {"DSC", "TGA"}:
        return "Not applicable"

    candidate_temperature = None
    rows = record.get("rows") or []
    summary = record.get("summary") or {}
    if analysis_type == "DSC":
        if rows:
            candidate_temperature = rows[0].get("peak_temperature")
        if candidate_temperature in (None, ""):
            candidate_temperature = summary.get("tg_midpoint")
    elif analysis_type == "TGA":
        if rows:
            candidate_temperature = rows[0].get("midpoint_temperature") or rows[0].get("onset_temperature")

    if candidate_temperature in (None, ""):
        return "No reference candidate recorded"

    try:
        candidate_temperature = float(candidate_temperature)
    except (TypeError, ValueError):
        return "Reference candidate could not be parsed"

    reference = find_nearest_reference(candidate_temperature, analysis_type=analysis_type)
    if reference is None:
        return f"No close reference match within 15.0 °C (candidate {candidate_temperature:.1f} °C)"

    delta = candidate_temperature - reference.temperature_c
    standard = f"; {reference.standard}" if reference.standard else ""
    return f"{reference.name} ({reference.temperature_c:.1f} °C, ΔT {delta:+.1f} °C{standard})"


def _xrd_caution_note(summary_payload: dict[str, Any]) -> str:
    match_status = str(summary_payload.get("match_status") or "").strip().lower()
    confidence_band = str(summary_payload.get("confidence_band") or "").strip().lower()
    caution_message = str(summary_payload.get("caution_message") or "").strip()
    best_candidate = _xrd_best_candidate_name(summary_payload)
    reason = str(summary_payload.get("top_candidate_reason_below_threshold") or "").strip()
    if match_status == "not_run":
        return caution_message or "Phase matching was not run because no reference library candidates were available."
    if match_status == "no_match":
        if caution_message:
            return caution_message
        if best_candidate != "N/A":
            if reason:
                return (
                    f"A best-ranked candidate ({best_candidate}) was identified, but it did not meet the configured "
                    f"qualitative acceptance threshold. Primary limiting factors: {reason}. Interpret as a screening "
                    "result rather than a confirmed identification."
                )
            return (
                f"A best-ranked candidate ({best_candidate}) was identified, but evidence remained insufficient for an "
                "accepted qualitative phase call. Interpret as a screening result rather than a confirmed identification."
            )
        return "No candidate exceeded the minimum match threshold; treat as qualitative caution outcome."
    if match_status == "matched" and confidence_band == "low":
        return caution_message or "An accepted candidate was retained, but confidence is low; review the peak-evidence table before interpretation."
    return "None"


def _domain_method_summary(record: dict) -> dict[str, str] | None:
    analysis_type = str(record.get("analysis_type") or "").upper()
    if analysis_type not in {"DSC", "TGA", "FTIR", "RAMAN", "XRD"}:
        return None

    processing = ensure_processing_payload(record.get("processing"), analysis_type=analysis_type) if record.get("processing") else {}
    method_context = processing.get("method_context") or {}

    if analysis_type == "DSC":
        summary = {
            "Template": processing.get("workflow_template_label") or processing.get("workflow_template") or "Not recorded",
            "Sign Convention": method_context.get("sign_convention_label") or processing.get("sign_convention") or "Not recorded",
            "Smoothing": _format_processing_step(_processing_step(processing, "smoothing")),
            "Baseline": _format_processing_step(_processing_step(processing, "baseline")),
            "Peak Analysis Context": _format_processing_step(_processing_step(processing, "peak_detection")),
            "Glass Transition Context": _format_processing_step(_processing_step(processing, "glass_transition")),
            "Reference Check": _reference_visibility(record),
        }
        return _table_payload(summary)

    if analysis_type == "TGA":
        summary = {
            "Template": processing.get("workflow_template_label") or processing.get("workflow_template") or "Not recorded",
            "Declared Unit Mode": method_context.get("tga_unit_mode_label") or "Not recorded",
            "Resolved Unit Mode": method_context.get("tga_unit_mode_resolved_label") or "Not recorded",
            "Smoothing": _format_processing_step(_processing_step(processing, "smoothing")),
            "Step Analysis Context": _format_processing_step(_processing_step(processing, "step_detection")),
            "Reference Check": _reference_visibility(record),
        }
        return _table_payload(summary)

    summary_payload = record.get("summary") or {}
    if analysis_type == "XRD":
        summary = {
            "Template": processing.get("workflow_template_label") or processing.get("workflow_template") or "Not recorded",
            "Axis Normalization": _format_processing_step(_processing_step(processing, "axis_normalization")),
            "Smoothing": _format_processing_step(_processing_step(processing, "smoothing")),
            "Baseline": _format_processing_step(_processing_step(processing, "baseline")),
            "Peak Detection": _format_processing_step(_processing_step(processing, "peak_detection")),
            "Match Metric": method_context.get("xrd_match_metric") or "Not recorded",
            "Tolerance (deg)": _format_value(method_context.get("xrd_match_tolerance_deg")),
            "Minimum Score": _format_value(method_context.get("xrd_match_minimum_score")),
        }
        return _table_payload(summary)

    summary = {
        "Template": processing.get("workflow_template_label") or processing.get("workflow_template") or "Not recorded",
        "Smoothing": _format_processing_step(_processing_step(processing, "smoothing")),
        "Baseline": _format_processing_step(_processing_step(processing, "baseline")),
        "Normalization": _format_processing_step(_processing_step(processing, "normalization")),
        "Peak Detection": _format_processing_step(_processing_step(processing, "peak_detection")),
        "Similarity Matching": _format_processing_step(_processing_step(processing, "similarity_matching")),
        "Match Status": summary_payload.get("match_status") or "Not recorded",
        "Confidence Band": summary_payload.get("confidence_band") or "Not recorded",
        "Caution Code": summary_payload.get("caution_code") or "None",
    }
    return _table_payload(summary)


def _generic_method_summary(processing: dict | None) -> dict[str, str]:
    processing = processing or {}
    return _table_payload(
        {
            "analysis_type": processing.get("analysis_type"),
            "workflow_template": processing.get("workflow_template"),
            "method": processing.get("method"),
        }
    )


def _processing_sections(processing: dict | None) -> list[tuple[str, dict[str, str]]]:
    processing = processing or {}
    if not processing:
        return []

    sections: list[tuple[str, dict[str, str]]] = []
    method_context = _compact_table_payload(processing.get("method_context"))
    if method_context:
        sections.append(("Method Context", method_context))

    signal_pipeline = processing.get("signal_pipeline")
    if not isinstance(signal_pipeline, dict) or not signal_pipeline:
        signal_pipeline = {
            key: processing.get(key)
            for key in ("smoothing", "baseline")
            if isinstance(processing.get(key), dict)
        }
    signal_payload = _compact_table_payload(signal_pipeline)
    if signal_payload:
        sections.append(("Signal Pipeline", signal_payload))

    analysis_steps = processing.get("analysis_steps")
    if not isinstance(analysis_steps, dict) or not analysis_steps:
        analysis_steps = {
            key: processing.get(key)
            for key in ("glass_transition", "peak_detection", "step_detection")
            if isinstance(processing.get(key), dict)
        }
    step_payload = _compact_table_payload(analysis_steps)
    if step_payload:
        sections.append(("Analysis Steps", step_payload))

    return sections


def _validation_sections(validation: dict | None) -> list[tuple[str, dict[str, str]]]:
    validation = validation or {}
    if not validation:
        return []

    sections: list[tuple[str, dict[str, str]]] = []
    summary = _table_payload(
        {
            "status": validation.get("status"),
            "issue_count": len(validation.get("issues") or []),
            "warning_count": len(validation.get("warnings") or []),
        }
    )
    if summary:
        sections.append(("Data Validation", summary))

    issue_payload = _table_payload({"issues": validation.get("issues")})
    if issue_payload:
        sections.append(("Validation Issues", issue_payload))

    warning_payload = _table_payload({"warnings": validation.get("warnings")})
    if warning_payload:
        sections.append(("Validation Warnings", warning_payload))

    checks_payload = _table_payload(validation.get("checks"))
    if checks_payload:
        sections.append(("Validation Checks", checks_payload))

    return sections


def _provenance_sections(provenance: dict | None) -> list[tuple[str, dict[str, str]]]:
    provenance = provenance or {}
    if not provenance:
        return []

    primary_keys = (
        "saved_at_utc",
        "dataset_key",
        "source_data_hash",
        "vendor",
        "instrument",
        "analyst_name",
        "app_version",
        "analysis_event_count",
    )
    sections: list[tuple[str, dict[str, str]]] = []

    summary = _table_payload({key: provenance.get(key) for key in primary_keys})
    if summary:
        sections.append(("Provenance", summary))

    remaining = {
        key: value
        for key, value in provenance.items()
        if key not in primary_keys and value not in (None, "", [], {})
    }
    context_payload = _table_payload(remaining)
    if context_payload:
        sections.append(("Provenance Context", context_payload))

    return sections


def _record_main_sections(record: dict) -> list[tuple[str, dict[str, Any]]]:
    context_sections = scientific_context_to_report_sections(record.get("scientific_context"))
    analysis_type = str(record.get("analysis_type") or "").upper()

    methodology_payload: dict[str, Any] = {}
    ordered_sections: list[tuple[str, dict[str, Any]]] = []

    for title, payload in context_sections:
        if title == "Methodology":
            if analysis_type != "XRD" and isinstance(payload, dict):
                methodology_payload.update(payload)
            continue
        if title == "Warnings and Limitations":
            continue
        ordered_sections.append((title, payload if isinstance(payload, dict) else {}))

    domain_summary = _domain_method_summary(record)
    if domain_summary is None:
        domain_summary = _generic_method_summary(record.get("processing"))
    methodology_payload.update(domain_summary or {})

    sections: list[tuple[str, dict[str, Any]]] = []
    if methodology_payload:
        sections.append(("Methodology", _table_payload(methodology_payload)))

    if analysis_type == "TGA":
        tga_narrative = build_tga_scientific_narrative(
            summary=record.get("summary") or {},
            rows=record.get("rows") or [],
            metadata=record.get("metadata") or {},
            validation=record.get("validation") or {},
        )
        if tga_narrative:
            ordered_sections = [
                section for section in ordered_sections
                if section[0] != "Primary Scientific Interpretation"
            ]
            insert_index = 1 if (ordered_sections and ordered_sections[0][0] == "Equations and Formulation") else 0
            ordered_sections.insert(
                insert_index,
                (
                    "Primary Scientific Interpretation",
                    {f"Observation {idx}": line for idx, line in enumerate(tga_narrative, start=1)},
                ),
            )

    section_priority = {
        "Equations and Formulation": 10,
        "Primary Scientific Interpretation": 20,
        "Evidence Supporting This Interpretation": 30,
        "Literature Comparison": 35,
        "Supporting References": 36,
        "Contradictory or Alternative References": 37,
        "Alternative Explanations": 40,
        "Uncertainty and Methodological Limits": 50,
        "Recommended Follow-Up Literature Checks": 55,
        "Recommended Follow-Up Experiments": 60,
        "Scientific Interpretation": 70,
        "Fit Quality": 80,
    }
    ordered_sections.extend(_literature_main_sections(record))
    ordered_sections.sort(key=lambda item: section_priority.get(item[0], 999))
    sections.extend(ordered_sections)

    grouped = condense_warning_limitations(
        record.get("scientific_context"),
        validation=record.get("validation"),
        max_data_warnings=3,
        max_method_limits=2,
    )
    warning_payload: dict[str, Any] = {}
    if grouped.get("data_completeness_warnings"):
        warning_payload["Data Completeness Warnings"] = grouped["data_completeness_warnings"]
    if grouped.get("methodological_limitations"):
        warning_payload["Methodological Limitations"] = grouped["methodological_limitations"]
    if warning_payload:
        sections.append(("Warnings and Limitations", warning_payload))
    return sections


def _citation_lookup(record: Mapping[str, Any]) -> dict[str, dict[str, Any]]:
    lookup: dict[str, dict[str, Any]] = {}
    for item in record.get("citations") or []:
        if not isinstance(item, Mapping):
            continue
        citation_id = normalize_report_text(item.get("citation_id") or "")
        if not citation_id:
            continue
        lookup[citation_id] = dict(item)
    return lookup


def _literature_demo_enabled() -> bool:
    token = (
        os.getenv("THERMOANALYZER_INCLUDE_DEMO_LITERATURE")
        or os.getenv("MATERIALSCOPE_INCLUDE_DEMO_LITERATURE")
        or ""
    )
    return token.strip().lower() in {"1", "true", "yes", "on"}


def _citation_is_fixture(citation: Mapping[str, Any]) -> bool:
    provenance = dict(citation.get("provenance") or {})
    provider_id = normalize_report_text(provenance.get("provider_id") or "").lower()
    result_source = normalize_report_text(provenance.get("result_source") or "").lower()
    provider_scope = [
        normalize_report_text(item).lower()
        for item in (provenance.get("provider_scope") or [])
        if normalize_report_text(item)
    ]
    return provider_id == "fixture_provider" or result_source == "fixture_search" or "fixture_provider" in provider_scope


def _literature_fixture_state(record: Mapping[str, Any]) -> dict[str, bool]:
    context = dict(record.get("literature_context") or {})
    provider_scope = [
        normalize_report_text(item).lower()
        for item in (context.get("provider_scope") or [])
        if normalize_report_text(item)
    ]
    citations = [dict(item) for item in (record.get("citations") or []) if isinstance(item, Mapping)]
    fixture_from_scope = "fixture_provider" in provider_scope
    fixture_citations = [citation for citation in citations if _citation_is_fixture(citation)]
    fixture_detected = fixture_from_scope or bool(fixture_citations)
    fixture_only = fixture_detected and (not citations or len(fixture_citations) == len(citations))
    return {
        "fixture_detected": fixture_detected,
        "fixture_only": fixture_only,
    }


def _comparison_is_fixture(
    comparison: Mapping[str, Any],
    *,
    citations_by_id: Mapping[str, Mapping[str, Any]],
) -> bool:
    citation_ids = [
        normalize_report_text(token)
        for token in (comparison.get("citation_ids") or [])
        if normalize_report_text(token)
    ]
    cited_rows = [dict(citations_by_id[citation_id]) for citation_id in citation_ids if citation_id in citations_by_id]
    if cited_rows:
        return all(_citation_is_fixture(citation) for citation in cited_rows)

    provider_id = normalize_report_text(comparison.get("provider_id") or "").lower()
    return provider_id == "fixture_provider"


def _citation_report_line(citation: Mapping[str, Any]) -> str:
    if _citation_is_fixture(citation):
        title = normalize_report_text(citation.get("title") or "Fixture citation metadata")
        access_class = normalize_report_text(citation.get("access_class") or "metadata_only")
        return normalize_report_text(
            f"Development/demo fixture citation only: {title}. Access class: {access_class}. "
            "This is not an authoritative bibliographic reference."
        )
    citation_text = _presentation_text(citation.get("citation_text") or citation.get("title") or "Citation not recorded")
    access_class = normalize_report_text(citation.get("access_class") or "metadata_only")
    link_markup = _citation_link_markup(citation)
    suffix = f" {link_markup}" if link_markup else ""
    return normalize_report_text(f"{citation_text} Access class: {access_class}.{suffix}")


def _citation_appendix_line(citation: Mapping[str, Any]) -> str:
    if _citation_is_fixture(citation):
        title = normalize_report_text(citation.get("title") or "Fixture citation metadata")
        access_class = normalize_report_text(citation.get("access_class") or "metadata_only")
        provenance = dict(citation.get("provenance") or {})
        provider_id = normalize_report_text(provenance.get("provider_id") or "fixture_provider")
        result_source = normalize_report_text(provenance.get("result_source") or "fixture_search")
        return normalize_report_text(
            f"Development/demo fixture citation only: {title}. Access class: {access_class}. "
            f"Provider: {provider_id}. Result source: {result_source}. "
            "DOI/URL text is intentionally omitted from authoritative citation display."
        )
    citation_text = _presentation_text(citation.get("citation_text") or citation.get("title") or "Citation not recorded")
    access_class = normalize_report_text(citation.get("access_class") or "metadata_only")
    license_note = normalize_report_text(citation.get("source_license_note") or "not recorded")
    provenance = dict(citation.get("provenance") or {})
    provider_id = normalize_report_text(provenance.get("provider_id") or "")
    result_source = normalize_report_text(provenance.get("result_source") or "")
    provider_scope = ", ".join(
        normalize_report_text(item)
        for item in (provenance.get("provider_scope") or [])
        if normalize_report_text(item)
    )
    request_ids = ", ".join(
        normalize_report_text(item)
        for item in ((provenance.get("provider_request_ids") or []) or [provenance.get("request_id")])
        if normalize_report_text(item)
    )
    segments = [
        citation_text,
        f"Access class: {access_class}.",
        f"License note: {license_note}.",
    ]
    if provider_id:
        segments.append(f"Provider: {provider_id}.")
    if result_source:
        segments.append(f"Result source: {result_source}.")
    if provider_scope:
        segments.append(f"Provider scope: {provider_scope}.")
    if request_ids:
        segments.append(f"Provider request IDs: {request_ids}.")
    link_markup = _citation_link_markup(citation)
    if link_markup:
        segments.append(link_markup)
    return normalize_report_text(" ".join(segments))


def _literature_main_sections(record: Mapping[str, Any]) -> list[tuple[str, dict[str, Any]]]:
    comparisons = [dict(item) for item in (record.get("literature_comparisons") or []) if isinstance(item, Mapping)]
    context = dict(record.get("literature_context") or {})
    analysis_type = str(record.get("analysis_type") or "").upper()
    all_citations_by_id = _citation_lookup(record)
    fixture_state = _literature_fixture_state(record)
    if fixture_state["fixture_only"]:
        return []
    citations_by_id = {
        citation_id: citation
        for citation_id, citation in all_citations_by_id.items()
        if not _citation_is_fixture(citation)
    }
    reportable_xrd_comparisons = [
        item
        for item in comparisons
        if not _comparison_is_fixture(item, citations_by_id=all_citations_by_id)
    ]
    if not comparisons and not citations_by_id:
        return []

    if analysis_type == "XRD" and (
        normalize_report_text(context.get("query_text") or "")
        or any(normalize_report_text(item.get("paper_title") or "") for item in reportable_xrd_comparisons)
    ):
        if not reportable_xrd_comparisons and not citations_by_id:
            return []
        sections: list[tuple[str, dict[str, Any]]] = []
        candidate_name = (
            normalize_report_text(context.get("candidate_display_name") or "")
            or normalize_report_text(context.get("candidate_name") or "")
            or normalize_report_text((record.get("summary") or {}).get("top_candidate_display_name_unicode") or "")
            or normalize_report_text((record.get("summary") or {}).get("top_candidate_name") or "")
            or "Not recorded"
        )
        candidate_context = {
            "Candidate": candidate_name,
            "Query Text": normalize_report_text(context.get("query_text") or "Not recorded"),
            "Accepted Match Status": normalize_report_text(context.get("match_status_snapshot") or (record.get("summary") or {}).get("match_status") or "Not recorded"),
            "Confidence Band": normalize_report_text(context.get("confidence_band_snapshot") or (record.get("summary") or {}).get("confidence_band") or "Not recorded"),
            "Authoritative Note": "Literature provides contextual evidence around the top-ranked candidate and does not override XRD screening.",
        }
        sections.append(("XRD Candidate Literature Check", candidate_context))

        relevant_payload: dict[str, Any] = {}
        alternative_payload: dict[str, Any] = {}
        for item in reportable_xrd_comparisons:
            title = normalize_report_text(item.get("paper_title") or item.get("claim_text") or f"Paper {len(relevant_payload) + len(alternative_payload) + 1}")
            journal = normalize_report_text(item.get("paper_journal") or "")
            year = normalize_report_text(item.get("paper_year") or "n.d.")
            link = _citation_link_markup({"doi": item.get("paper_doi"), "url": item.get("paper_url")})
            posture = normalize_report_text(item.get("validation_posture") or "non_validating")
            provider = normalize_report_text(item.get("provider_id") or "not_recorded")
            access_class = normalize_report_text(item.get("access_class") or "metadata_only")
            note = normalize_report_text(item.get("comparison_note") or item.get("rationale") or "No comparison note recorded.")
            line = f"{year}. {journal}. provider={provider}; access={access_class}; posture={posture}. {note}"
            if link:
                line = f"{line} {link}"
            if posture == "related_support":
                relevant_payload[title] = normalize_report_text(line)
            else:
                alternative_payload[title] = normalize_report_text(line)

        if not relevant_payload and reportable_xrd_comparisons:
            relevant_payload["Note"] = "No paper met the threshold for candidate-related support; literature remains contextual and non-definitive."
        sections.append(("Relevant Papers", relevant_payload))

        if not alternative_payload and reportable_xrd_comparisons:
            alternative_payload["Note"] = "No alternative or non-validating papers were retained beyond the contextual candidate check."
        sections.append(("Alternative or Non-Validating Papers", alternative_payload))

        follow_up_checks: dict[str, Any] = {}
        if not reportable_xrd_comparisons:
            follow_up_checks["Check 1"] = "No bibliographic papers were retained for the top-ranked candidate; consider broadening candidate labels or searching additional metadata providers."
        if normalize_report_text(context.get("match_status_snapshot") or "").lower() == "no_match":
            follow_up_checks["Check 2"] = "The accepted XRD outcome remains no_match; literature does not validate a phase call and confirmatory experiments may still be required."
        if context.get("metadata_only_evidence"):
            follow_up_checks["Check 3"] = "The retained paper set is metadata- or abstract-heavy; broader open-access evidence could refine candidate-level context."
        if follow_up_checks:
            sections.append(("Recommended Follow-Up Literature Checks", follow_up_checks))
        return sections

    sections: list[tuple[str, dict[str, Any]]] = []
    comparison_payload: dict[str, Any] = {}
    if analysis_type in {"DSC", "DTA", "TGA"} and normalize_report_text(context.get("query_text") or ""):
        thermal_context = {
            "Search Focus": _presentation_text(context.get("query_display_title") or context.get("candidate_display_name") or context.get("candidate_name") or "Not recorded"),
            "Search Mode": normalize_report_text(context.get("query_display_mode") or "Thermal / interpretation"),
            "Query Text": _presentation_text(context.get("query_text") or "Not recorded"),
            "Authoritative Note": "Literature provides context around the recorded thermal interpretation and does not validate or override the current result.",
        }
        if normalize_report_text(context.get("query_rationale") or ""):
            thermal_context["Query Rationale"] = _presentation_text(context.get("query_rationale"))
        if context.get("low_specificity_retrieval"):
            thermal_context["Retrieval Note"] = (
                "Real literature results were found, but the retained set is low-specificity and mostly metadata/abstract-level; direct validation remains unavailable."
            )
        evidence_specificity = normalize_report_text(context.get("evidence_specificity_summary") or "").lower()
        if evidence_specificity == "abstract_backed":
            thermal_context["Evidence Basis"] = (
                "At least one retained source includes accessible abstract-level evidence; interpretation remains non-validating but is better grounded than metadata-only retrieval."
            )
        elif evidence_specificity == "mixed_metadata_and_abstract":
            thermal_context["Evidence Basis"] = (
                "The retained source set mixes metadata-only and accessible abstract evidence; interpretation remains cautious, but it is not purely metadata-based."
            )
        elif evidence_specificity == "oa_backed":
            thermal_context["Evidence Basis"] = (
                "At least one retained source includes accessible open text; the interpretation remains cautious and non-validating."
            )
        sections.append(("Thermal Literature Search Summary", thermal_context))

    for item in comparisons:
        claim_id = normalize_report_text(item.get("claim_id") or f"C{len(comparison_payload) + 1}")
        label = normalize_report_text(item.get("support_label") or "related_but_inconclusive")
        confidence = normalize_report_text(item.get("confidence") or "low")
        rationale = _presentation_text(item.get("rationale") or "No literature rationale recorded.")
        citation_ids = [
            normalize_report_text(token)
            for token in (item.get("citation_ids") or [])
            if normalize_report_text(token) and normalize_report_text(token) in citations_by_id
        ]
        citation_note = f" Citations: {', '.join(citation_ids)}." if citation_ids else " No accessible citation was retained for this claim."
        comparison_payload[f"{claim_id} ({label}, {confidence})"] = normalize_report_text(f"{rationale}{citation_note}")
    if comparison_payload:
        sections.append(("Literature Comparison", comparison_payload))

    supporting_ids, alternative_ids = partition_reference_ids(
        comparisons,
        citations_by_id=citations_by_id,
        context={**context, "analysis_type": analysis_type},
        analysis_type=analysis_type,
    )

    supporting_payload = {
        citation_id: _citation_report_line(citations_by_id[citation_id])
        for citation_id in supporting_ids
        if citation_id in citations_by_id
    }
    if not supporting_payload and comparisons:
        supporting_payload["Note"] = (
            "No supporting accessible references were retained for this run; evidence is limited and remains non-definitive."
        )
    if supporting_payload:
        sections.append(
            ("Relevant References", supporting_payload)
            if analysis_type in {"DSC", "DTA", "TGA"}
            else ("Supporting References", supporting_payload)
        )

    alternative_payload = {
        citation_id: _citation_report_line(citations_by_id[citation_id])
        for citation_id in alternative_ids
        if citation_id in citations_by_id
    }
    if not alternative_payload and comparisons:
        alternative_payload["Note"] = (
            "No contradictory accessible references were retained, but the current comparison still remains qualitative and cautionary."
        )
    if alternative_payload:
        sections.append(
            ("Alternative or Non-Validating References", alternative_payload)
            if analysis_type in {"DSC", "DTA", "TGA"}
            else ("Contradictory or Alternative References", alternative_payload)
        )

    follow_up_checks: dict[str, Any] = {}
    if any(not (item.get("citation_ids") or []) for item in comparisons):
        follow_up_checks["Check 1"] = (
            "At least one claim remained citation-limited; additional confirmatory experiments may be required before stronger literature alignment is inferred."
        )
    if any(str(item.get("confidence") or "").lower() == "low" for item in comparisons):
        follow_up_checks["Check 2"] = (
            "Low-confidence literature outcomes should be treated as screening context only, not as confirmation."
        )
    citation_access_classes = {
        normalize_report_text(item.get("access_class") or "")
        for item in citations_by_id.values()
    }
    if context.get("metadata_only_evidence") or citation_access_classes == {"metadata_only"}:
        follow_up_checks["Check 3"] = (
            "Some literature reasoning relies on metadata or abstract-level evidence only; broader open-access or user-provided documents could refine the comparison."
        )
    if context.get("low_specificity_retrieval"):
        follow_up_checks["Check 3b"] = (
            "The retained real-literature set is low-specificity and may reflect neighboring materials/process papers rather than direct thermal validation."
        )
    evidence_specificity = normalize_report_text(context.get("evidence_specificity_summary") or "").lower()
    if evidence_specificity in {"abstract_backed", "mixed_metadata_and_abstract"}:
        follow_up_checks["Check 3c"] = (
            "At least one retained source includes accessible abstract-level evidence; interpretation is better grounded than metadata-only retrieval but remains non-validating."
        )
    elif evidence_specificity == "oa_backed":
        follow_up_checks["Check 3c"] = (
            "At least one retained source includes accessible open text; interpretation remains cautious and does not override the current result."
        )
    if context.get("restricted_content_used") is False:
        follow_up_checks["Check 4"] = (
            "Closed-access full text was intentionally excluded from reasoning; the system remains legal-safe by design."
        )
    if follow_up_checks:
        sections.append(("Recommended Follow-Up Literature Checks", follow_up_checks))

    return sections


def _literature_appendix_sections(record: Mapping[str, Any]) -> list[tuple[str, dict[str, str]]]:
    sections: list[tuple[str, dict[str, str]]] = []
    fixture_state = _literature_fixture_state(record)
    if fixture_state["fixture_only"] and not _literature_demo_enabled():
        return [
            (
                "Development / Demo Literature Output",
                {
                    "Status": (
                        "Fixture/demo-only literature output was excluded from the report by default because it is not a real bibliographic source."
                    )
                },
            )
        ]

    context = dict(record.get("literature_context") or {})
    if context:
        context_payload = _table_payload(context)
        candidate_label = (
            normalize_report_text(context.get("candidate_display_name") or "")
            or normalize_report_text(context.get("candidate_name") or "")
            or normalize_report_text((record.get("summary") or {}).get("top_candidate_display_name_unicode") or "")
            or normalize_report_text((record.get("summary") or {}).get("top_candidate_name") or "")
        )
        if candidate_label:
            context_payload["Candidate Label"] = candidate_label
        reason = normalize_report_text(context.get("no_results_reason") or "").lower()
        query_status = normalize_report_text(context.get("provider_query_status") or "").lower()
        real_literature_available = bool(context.get("real_literature_available"))
        source_count = int(context.get("source_count") or 0)
        citation_count = int(context.get("citation_count") or 0)
        if (
            str(record.get("analysis_type") or "").upper() == "XRD"
            and not real_literature_available
            and source_count == 0
            and citation_count == 0
        ):
            if reason == "provider_unavailable" or query_status == "provider_unavailable":
                context_payload["Real Literature Search Outcome"] = (
                    "The configured real provider could not return usable bibliographic results for this candidate in the current run."
                )
            elif reason == "request_failed" or query_status == "request_failed":
                context_payload["Real Literature Search Outcome"] = (
                    "A request reached the configured real provider, but it did not return a usable bibliographic response for this candidate-centered query."
                )
            elif reason == "not_configured" or query_status == "not_configured":
                context_payload["Real Literature Search Outcome"] = (
                    "A live real-provider client was not configured for this environment, so no bibliographic papers were retrieved for the candidate-centered query."
                )
            elif reason == "query_too_narrow" or query_status == "query_too_narrow":
                context_payload["Real Literature Search Outcome"] = (
                    "The candidate-centered XRD query was too narrow to retrieve usable bibliographic metadata in this run; a clearer candidate label or formula may be required."
                )
            else:
                context_payload["Real Literature Search Outcome"] = (
                    "A real provider query completed but did not return suitable bibliographic results for the candidate-centered XRD search. "
                    "This does not imply that the phase is absent from the literature."
                )
            context_payload["Authoritative Note"] = (
                "XRD accepted match status remains authoritative; literature absence does not validate or invalidate the phase call."
            )
        title = "Literature Comparison Context"
        if fixture_state["fixture_only"]:
            title = "Development / Demo Literature Output"
        sections.append((title, context_payload))

    comparisons_payload = {}
    for item in record.get("literature_comparisons") or []:
        if not isinstance(item, Mapping):
            continue
        claim_id = normalize_report_text(item.get("claim_id") or f"C{len(comparisons_payload) + 1}")
        label = normalize_report_text(item.get("support_label") or "related_but_inconclusive")
        confidence = normalize_report_text(item.get("confidence") or "low")
        rationale = normalize_report_text(item.get("rationale") or "No literature rationale recorded.")
        citations = ", ".join(
            normalize_report_text(token)
            for token in (item.get("citation_ids") or [])
            if normalize_report_text(token)
        )
        comparisons_payload[f"{claim_id} ({label})"] = normalize_report_text(
            f"confidence={confidence}; citations={citations or 'none'}; {rationale}"
        )
    if comparisons_payload:
        sections.append(
            (
                "Demo Literature Comparison Outcomes" if fixture_state["fixture_only"] else "Literature Comparison Outcomes",
                comparisons_payload,
            )
        )

    citations_payload = {}
    for citation_id, citation in _citation_lookup(record).items():
        if _citation_is_fixture(citation) and not _literature_demo_enabled():
            continue
        citations_payload[citation_id] = _citation_appendix_line(citation)
    if citations_payload:
        sections.append(
            (
                "Demo Literature Citations" if fixture_state["fixture_only"] else "Literature Citations",
                citations_payload,
            )
        )
    return sections


def _xrd_appendix_summary_sections(record: dict) -> list[tuple[str, dict[str, str]]]:
    if str(record.get("analysis_type") or "").upper() != "XRD":
        return []
    summary = record.get("summary") or {}
    sections: list[tuple[str, dict[str, str]]] = []

    library_payload = _table_payload(
        {
            "Reference Candidate Count": summary.get("reference_candidate_count"),
            "Library Provider": summary.get("library_provider"),
            "Library Package": summary.get("library_package"),
            "Library Version": summary.get("library_version"),
            "Library Sync Mode": summary.get("library_sync_mode"),
            "Library Cache Status": summary.get("library_cache_status"),
            "Library Access Mode": summary.get("library_access_mode"),
            "Library Request ID": summary.get("library_request_id"),
            "Library Result Source": summary.get("library_result_source"),
            "Library Provider Scope": summary.get("library_provider_scope"),
            "Library Offline Limited Mode": summary.get("library_offline_limited_mode"),
            "Top Candidate Provider": summary.get("top_candidate_provider"),
            "Top Candidate Package": summary.get("top_candidate_package"),
            "Top Candidate Version": summary.get("top_candidate_version"),
        }
    )
    if library_payload:
        sections.append(("XRD Library and Access Context", library_payload))

    provenance_payload = _table_payload(
        {
            "Provider Candidate Counts": summary.get("xrd_provider_candidate_counts"),
            "XRD Coverage Tier": summary.get("xrd_coverage_tier"),
            "XRD Coverage Warning": summary.get("xrd_coverage_warning_message"),
            "XRD Provenance State": summary.get("xrd_provenance_state"),
            "XRD Provenance Warning": summary.get("xrd_provenance_warning"),
            "Match Tolerance (deg)": summary.get("match_tolerance_deg"),
            "Weighted Overlap Score": summary.get("top_candidate_weighted_overlap_score"),
            "Mean Delta Position": summary.get("top_candidate_mean_delta_position"),
            "Unmatched Major Peak Count": summary.get("top_candidate_unmatched_major_peak_count"),
            "Candidate Count": summary.get("candidate_count"),
        }
    )
    if provenance_payload:
        sections.append(("XRD Match and Provenance Context", provenance_payload))
    return sections


def _record_appendix_matrix_sections(record: dict) -> list[tuple[str, list[str], list[list[str]]]]:
    if str(record.get("analysis_type") or "").upper() != "XRD":
        return []
    rows = record.get("rows") or []
    if not rows:
        return []

    matrix: list[list[str]] = []
    for row in rows[:5]:
        evidence = dict(row.get("evidence") or {})
        matrix.append(
            [
                _format_value(row.get("rank")),
                xrd_candidate_display_name(row, target="unicode") or _format_value(row.get("candidate_name")),
                _format_number(row.get("normalized_score"), digits=3),
                _format_value(row.get("confidence_band")),
                _format_value(evidence.get("shared_peak_count")),
                _format_number(evidence.get("coverage_ratio"), digits=3),
                _format_number(evidence.get("weighted_overlap_score"), digits=3),
                _format_number(evidence.get("mean_delta_position"), digits=3),
                _format_value(evidence.get("unmatched_major_peak_count")),
                _format_value(len(evidence.get("matched_peak_pairs") or [])),
                _format_value(len(evidence.get("unmatched_observed_peaks") or [])),
                _format_value(len(evidence.get("unmatched_reference_peaks") or [])),
            ]
        )
    if not matrix:
        return []
    return [
        (
            "Candidate Evidence Summary",
            [
                "Rank",
                "Candidate",
                "Score",
                "Status",
                "Shared Peaks",
                "Coverage",
                "Weighted Overlap",
                "Mean ΔPos",
                "Unmatched Major",
                "Matched Pairs",
                "Unmatched Obs",
                "Unmatched Ref",
            ],
            matrix,
        )
    ]


def _record_appendix_sections(record: dict) -> list[tuple[str, dict[str, str]]]:
    sections: list[tuple[str, dict[str, str]]] = []
    sections.extend(_processing_sections(record.get("processing")))
    sections.extend(_xrd_appendix_summary_sections(record))
    sections.extend(_validation_sections(record.get("validation")))
    sections.extend(_provenance_sections(record.get("provenance")))
    metadata_payload = _appendix_dataset_metadata(record.get("metadata") or {})
    if metadata_payload:
        sections.append(("Dataset Metadata (Technical)", metadata_payload))
    review_payload = _table_payload(record.get("review"))
    if review_payload:
        sections.append(("Internal Review Context", review_payload))
    sections.extend(_literature_appendix_sections(record))

    context = normalize_scientific_context(record.get("scientific_context"))
    context_warnings = {f"Warning {idx}": normalize_report_text(item) for idx, item in enumerate(context.get("warnings") or [], start=1)}
    if context_warnings:
        sections.append(("Scientific Context Warnings (Extended)", context_warnings))
    context_limits = {f"Limitation {idx}": normalize_report_text(item) for idx, item in enumerate(context.get("limitations") or [], start=1)}
    if context_limits:
        sections.append(("Scientific Context Limitations (Extended)", context_limits))
    return sections


def _xrd_reference_dossiers(record: dict) -> list[dict[str, Any]]:
    if str(record.get("analysis_type") or "").upper() != "XRD":
        return []
    payload = (record.get("report_payload") or {}).get("xrd_reference_dossiers") or []
    return [dict(item) for item in payload if isinstance(item, Mapping)]


def _xrd_reference_dossier_overview_payload(dossier: Mapping[str, Any]) -> dict[str, str]:
    overview = dict(dossier.get("candidate_overview") or {})
    return _table_payload(
        {
            "rank": dossier.get("rank"),
            "candidate": overview.get("display_name_unicode") or overview.get("display_name"),
            "scientific_formula": overview.get("formula_unicode") or overview.get("formula"),
            "raw_label": overview.get("raw_label"),
            "candidate_id": overview.get("candidate_id"),
            "source_id": overview.get("source_id"),
            "provider": overview.get("provider"),
            "package": overview.get("package"),
            "package_version": overview.get("package_version"),
            "confidence_band": overview.get("confidence_band"),
            "match_status": overview.get("match_status"),
            "candidate_score": overview.get("candidate_score"),
            "canonical_material_key": overview.get("canonical_material_key"),
        }
    )


def _xrd_reference_dossier_match_payload(dossier: Mapping[str, Any]) -> dict[str, str]:
    evidence = dict(dossier.get("match_evidence") or {})
    return _table_payload(
        {
            "shared_peak_count": evidence.get("shared_peak_count"),
            "weighted_overlap_score": evidence.get("weighted_overlap_score"),
            "coverage_ratio": evidence.get("coverage_ratio"),
            "mean_delta_position": evidence.get("mean_delta_position"),
            "unmatched_major_peak_count": evidence.get("unmatched_major_peak_count"),
            "matched_peak_pair_count": evidence.get("matched_peak_pair_count"),
            "unmatched_observed_count": evidence.get("unmatched_observed_count"),
            "unmatched_reference_count": evidence.get("unmatched_reference_count"),
            "reason_below_threshold": evidence.get("reason_below_threshold"),
            "caution_note": evidence.get("caution_note"),
        }
    )


def _xrd_reference_dossier_metadata_payload(dossier: Mapping[str, Any]) -> dict[str, str]:
    metadata = dict(dossier.get("reference_metadata") or {})
    return _table_payload(
        {
            "provider": metadata.get("provider"),
            "package_id": metadata.get("package_id"),
            "package_version": metadata.get("package_version"),
            "provider_dataset_version": metadata.get("provider_dataset_version"),
            "hosted_dataset_version": metadata.get("hosted_dataset_version"),
            "hosted_published_at": metadata.get("hosted_published_at"),
            "published_at": metadata.get("published_at"),
            "generated_at": metadata.get("generated_at"),
            "last_updated": metadata.get("last_updated"),
            "canonical_material_key": metadata.get("canonical_material_key"),
            "space_group": metadata.get("space_group"),
            "symmetry": metadata.get("symmetry"),
            "formula": metadata.get("formula_unicode") or metadata.get("formula_pretty") or metadata.get("formula"),
            "phase_name": metadata.get("phase_name"),
            "display_name": metadata.get("display_name_unicode") or metadata.get("display_name"),
            "source_url": metadata.get("source_url"),
            "provider_url": metadata.get("provider_url"),
            "attribution": metadata.get("attribution"),
        }
    )


def _xrd_reference_peak_display_payload(dossier: Mapping[str, Any]) -> dict[str, str]:
    reference_peaks = dict(dossier.get("reference_peaks") or {})
    displayed = int(reference_peaks.get("displayed_peak_count") or 0)
    total = int(reference_peaks.get("total_peak_count") or 0)
    truncated = int(reference_peaks.get("truncated_count") or 0)
    selection_policy = str(reference_peaks.get("selection_policy") or "").strip()
    readable_policy = ""
    if selection_policy == "matched_and_major_then_fill_to_top_20_by_intensity":
        readable_policy = "Matched and major peaks were prioritized, then remaining slots were filled by descending intensity."
    return _table_payload(
        {
            "displayed_reference_peaks": displayed,
            "total_reference_peaks": total,
            "truncated_reference_peaks": truncated,
            "selection_policy": readable_policy or selection_policy,
        }
    )


def _xrd_reference_peak_matrix(dossier: Mapping[str, Any]) -> tuple[list[str], list[list[str]], list[str]]:
    reference_peaks = dict(dossier.get("reference_peaks") or {})
    headers = ["Peak #", "2θ", "d-spacing", "Relative Intensity", "Matched", "Major"]
    rows = [
        [
            _format_value(item.get("peak_number")),
            _format_number(item.get("position"), digits=3),
            _format_number(item.get("d_spacing"), digits=4),
            _format_number(item.get("relative_intensity"), digits=3),
            _format_value(item.get("matched")),
            _format_value(item.get("major")),
        ]
        for item in (reference_peaks.get("display_rows") or [])
        if isinstance(item, Mapping)
    ]
    displayed = int(reference_peaks.get("displayed_peak_count") or len(rows) or 0)
    total = int(reference_peaks.get("total_peak_count") or len(rows) or 0)
    notes: list[str] = []
    if total > displayed > 0:
        notes.append(f"Showing {displayed} of {total} reference peaks.")
        notes.append("Remaining peaks omitted from visible table by display policy.")
    elif total and not rows:
        notes.append(f"{total} reference peaks were available, but none were selected for display.")
    policy = _xrd_reference_peak_display_payload(dossier).get("Selection Policy")
    if policy:
        notes.append(policy)
    return headers, rows, notes


def _xrd_structure_visual_payload(dossier: Mapping[str, Any]) -> dict[str, str]:
    structure = dict(dossier.get("structure_payload") or {})
    metadata = dict(dossier.get("reference_metadata") or {})
    source_assets = [dict(item) for item in (dossier.get("source_assets") or []) if isinstance(item, Mapping)]
    availability = normalize_report_text(structure.get("availability") or "none").replace("_", " ")
    payload = _table_payload(
        {
            "availability_state": availability,
            "formula": structure.get("formula_unicode") or metadata.get("formula_unicode") or structure.get("formula") or metadata.get("formula_pretty") or metadata.get("formula"),
            "formula_raw": structure.get("formula") or metadata.get("formula"),
            "space_group": structure.get("space_group") or metadata.get("space_group"),
            "symmetry": structure.get("symmetry") or metadata.get("symmetry"),
            "source_url": structure.get("source_url") or metadata.get("source_url"),
            "provider_url": structure.get("provider_url") or metadata.get("provider_url"),
            "linked_source_count": structure.get("source_asset_count") or len(source_assets),
            "embeddable_visual_asset_count": structure.get("rendered_asset_count"),
            "notes": structure.get("notes") or metadata.get("attribution"),
        }
    )
    if not source_assets:
        payload["Fallback Note"] = normalize_report_text(XRD_NO_VISUAL_ASSET_NOTE)
    return payload


def _xrd_asset_label(dossier: Mapping[str, Any], asset: Mapping[str, Any], *, index: int) -> str:
    label = normalize_report_text(asset.get("label") or f"Asset {index}")
    provider = normalize_report_text(
        (dossier.get("candidate_overview") or {}).get("provider")
        or (dossier.get("reference_metadata") or {}).get("provider")
        or ""
    )
    if provider and label in {"Source Reference", "Provider Reference", "Reference Image", "Crystal Structure Asset"}:
        return normalize_report_text(f"{provider} {label}")
    return label


def _xrd_source_asset_matrix(dossier: Mapping[str, Any]) -> tuple[list[str], list[list[str]]]:
    rows: list[list[str]] = []
    for index, asset in enumerate((dossier.get("source_assets") or []), start=1):
        if not isinstance(asset, Mapping):
            continue
        location = normalize_report_text(
            asset.get("url")
            or asset.get("artifact_key")
            or "Not recorded"
        )
        rows.append(
            [
                _xrd_asset_label(dossier, asset, index=index),
                _format_value(asset.get("kind")).replace("_", " "),
                _format_value(asset.get("available")),
                location,
            ]
        )
    return ["Label", "Type", "Available", "Link / Asset"], rows


def _xrd_source_asset_lines(dossier: Mapping[str, Any]) -> list[str]:
    lines: list[str] = []
    for index, asset in enumerate((dossier.get("source_assets") or []), start=1):
        if not isinstance(asset, Mapping):
            continue
        location = str(asset.get("url") or asset.get("artifact_key") or "").strip()
        if not location:
            continue
        lines.append(f"{_xrd_asset_label(dossier, asset, index=index)}: {location}")
    return lines


def _xrd_provenance_payload(dossier: Mapping[str, Any]) -> dict[str, str]:
    provenance = dict(dossier.get("provenance") or {})
    return _table_payload(
        {
            "provider": provenance.get("provider"),
            "package": provenance.get("package"),
            "package_version": provenance.get("package_version"),
            "library_request_id": provenance.get("library_request_id"),
            "candidate_id": provenance.get("candidate_id"),
            "source_id": provenance.get("source_id"),
            "raw_label": provenance.get("raw_label"),
            "attribution": provenance.get("attribution"),
        }
    )


def _xrd_renderable_assets(dossier: Mapping[str, Any], figures: dict | None) -> list[tuple[str, bytes]]:
    rendered: list[tuple[str, bytes]] = []
    if not isinstance(figures, Mapping):
        return rendered
    for asset in dossier.get("source_assets") or []:
        if not isinstance(asset, Mapping):
            continue
        artifact_key = str(asset.get("artifact_key") or "").strip()
        if not artifact_key:
            continue
        png_bytes = figures.get(artifact_key)
        if isinstance(png_bytes, (bytes, bytearray)):
            rendered.append((str(asset.get("label") or artifact_key), bytes(png_bytes)))
    return rendered


def _render_docx_matrix_table(doc: Document, headers: list[str], rows: list[list[Any]], *, heading: str | None = None, heading_level: int = 3) -> None:
    if not headers or not rows:
        return
    portrait_width = _docx_section_text_width_inches(doc.sections[-1]) * 72.0
    layout = _choose_portrait_or_landscape_table_layout(headers, rows, portrait_width=portrait_width)
    if layout == "landscape":
        landscape_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        _set_docx_section_orientation(landscape_section, landscape=True)
    if heading:
        doc.add_paragraph(normalize_report_text(heading), style=f"Heading {heading_level}")
    _add_results_table(doc, headers, rows)
    doc.add_paragraph()
    if layout == "landscape":
        portrait_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        _set_docx_section_orientation(portrait_section, landscape=False)


def _render_xrd_reference_dossier_docx(doc: Document, dossier: Mapping[str, Any], *, figures: dict | None) -> None:
    rank = _format_value(dossier.get("rank"))
    doc.add_paragraph(normalize_report_text(f"Candidate Reference Dossier — Rank #{rank}"), style="Heading 3")

    doc.add_paragraph("Candidate Overview", style="Heading 4")
    _add_key_value_table(doc, _xrd_reference_dossier_overview_payload(dossier))
    doc.add_paragraph()

    doc.add_paragraph("Match Evidence Summary", style="Heading 4")
    _add_key_value_table(doc, _xrd_reference_dossier_match_payload(dossier))
    doc.add_paragraph()

    doc.add_paragraph("Reference Metadata", style="Heading 4")
    _add_key_value_table(doc, _xrd_reference_dossier_metadata_payload(dossier))
    doc.add_paragraph()

    doc.add_paragraph("Reference Peaks", style="Heading 4")
    _add_key_value_table(doc, _xrd_reference_peak_display_payload(dossier))
    doc.add_paragraph()
    headers, rows, notes = _xrd_reference_peak_matrix(dossier)
    if rows:
        _render_docx_matrix_table(doc, headers, rows)
    for note in notes:
        doc.add_paragraph(normalize_report_text(note))
    if rows or notes:
        doc.add_paragraph()

    doc.add_paragraph("Structure / Visual Evidence", style="Heading 4")
    _add_key_value_table(doc, _xrd_structure_visual_payload(dossier))
    doc.add_paragraph()
    asset_headers, asset_rows = _xrd_source_asset_matrix(dossier)
    if asset_rows:
        _render_docx_matrix_table(doc, asset_headers, asset_rows, heading="Linked Source and Provider Assets", heading_level=4)
    for line in _xrd_source_asset_lines(dossier):
        doc.add_paragraph(normalize_report_text(line))
    if asset_rows or _xrd_source_asset_lines(dossier):
        doc.add_paragraph()
    rendered_assets = _xrd_renderable_assets(dossier, figures)
    if rendered_assets:
        for label, png_bytes in rendered_assets:
            try:
                doc.add_picture(io.BytesIO(png_bytes), width=Inches(4.8))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph(normalize_report_text(label))
            except Exception:
                doc.add_paragraph(normalize_report_text(f"{label}: asset could not be embedded."))
        doc.add_paragraph()
    else:
        doc.add_paragraph(normalize_report_text(XRD_NO_VISUAL_ASSET_NOTE))
        doc.add_paragraph()

    doc.add_paragraph("Provenance / Attribution", style="Heading 4")
    _add_key_value_table(doc, _xrd_provenance_payload(dossier))
    doc.add_paragraph()


def _render_xrd_reference_dossier_pdf(
    story,
    dossier: Mapping[str, Any],
    *,
    figures: dict | None,
    add_heading,
    add_kv_table,
    add_matrix_table,
    ensure_template,
    portrait_width: float,
    landscape_width: float,
    body_style,
    small_style,
    caption_style,
    paragraph_factory,
    image_factory,
    spacer_factory,
) -> None:
    rank = _format_value(dossier.get("rank"))
    ensure_template("Portrait")
    add_heading(f"Candidate Reference Dossier — Rank #{rank}", level=3)

    add_heading("Candidate Overview", level=3)
    add_kv_table(_xrd_reference_dossier_overview_payload(dossier), width=portrait_width, compact=True)
    story.append(spacer_factory(1, 4))

    add_heading("Match Evidence Summary", level=3)
    add_kv_table(_xrd_reference_dossier_match_payload(dossier), width=portrait_width, compact=True)
    story.append(spacer_factory(1, 4))

    add_heading("Reference Metadata", level=3)
    add_kv_table(_xrd_reference_dossier_metadata_payload(dossier), width=portrait_width, compact=True)
    story.append(spacer_factory(1, 4))

    add_heading("Reference Peaks", level=3)
    add_kv_table(_xrd_reference_peak_display_payload(dossier), width=portrait_width, compact=True)
    story.append(spacer_factory(1, 4))
    headers, rows, notes = _xrd_reference_peak_matrix(dossier)
    if rows:
        matrix_layout = _choose_portrait_or_landscape_table_layout(headers, rows, portrait_width=portrait_width)
        ensure_template("Landscape" if matrix_layout == "landscape" else "Portrait")
        add_matrix_table(headers, rows, width=landscape_width if matrix_layout == "landscape" else portrait_width, compact=True)
    for note in notes:
        story.append(spacer_factory(1, 2))
        story.append(paragraph_factory(normalize_report_text(note), small_style))
    if rows or notes:
        story.append(spacer_factory(1, 4))

    ensure_template("Portrait")
    add_heading("Structure / Visual Evidence", level=3)
    add_kv_table(_xrd_structure_visual_payload(dossier), width=portrait_width, compact=True)
    story.append(spacer_factory(1, 3))
    asset_headers, asset_rows = _xrd_source_asset_matrix(dossier)
    asset_lines = _xrd_source_asset_lines(dossier)
    if asset_rows:
        add_heading("Linked Source and Provider Assets", level=3)
        add_matrix_table(asset_headers, asset_rows, width=portrait_width, compact=True)
        story.append(spacer_factory(1, 4))
    for line in asset_lines:
        story.append(paragraph_factory(normalize_report_text(line), body_style))
    if asset_lines:
        story.append(spacer_factory(1, 4))
    rendered_assets = _xrd_renderable_assets(dossier, figures)
    if rendered_assets:
        for label, png_bytes in rendered_assets:
            try:
                image = image_factory(io.BytesIO(png_bytes))
                image._restrictSize(portrait_width * 0.72, 180)
                story.append(spacer_factory(1, 3))
                story.append(image)
                story.append(paragraph_factory(normalize_report_text(label), caption_style))
            except Exception:
                story.append(paragraph_factory(normalize_report_text(f"{label}: asset could not be embedded."), body_style))
        story.append(spacer_factory(1, 4))
    else:
        story.append(spacer_factory(1, 2))
        story.append(paragraph_factory(normalize_report_text(XRD_NO_VISUAL_ASSET_NOTE), body_style))
        story.append(spacer_factory(1, 4))

    add_heading("Provenance / Attribution", level=3)
    add_kv_table(_xrd_provenance_payload(dossier), width=portrait_width, compact=True)
    story.append(spacer_factory(1, 6))


def _render_record_mapping(doc: Document, title: str, payload: dict | None) -> None:
    payload = payload or {}
    if not payload:
        return

    doc.add_paragraph(normalize_report_text(title), style="Heading 3")
    if title in _BULLET_SECTION_TITLES:
        for key, value in payload.items():
            bullet = normalize_report_text(value)
            if title not in {"Scientific Interpretation", "Alternative Explanations", "Recommended Follow-Up Experiments"}:
                bullet = normalize_report_text(f"{key}: {value}")
            paragraph = doc.add_paragraph(style="List Bullet")
            _append_docx_text_with_links(paragraph, bullet)
        doc.add_paragraph()
        return

    if title == "Warnings and Limitations":
        for group, values in payload.items():
            doc.add_paragraph(normalize_report_text(group), style="Heading 4")
            if isinstance(values, list):
                for item in values:
                    paragraph = doc.add_paragraph(style="List Bullet")
                    _append_docx_text_with_links(paragraph, normalize_report_text(item))
            else:
                paragraph = doc.add_paragraph(style="List Bullet")
                _append_docx_text_with_links(paragraph, normalize_report_text(values))
        doc.add_paragraph()
        return

    _add_key_value_table(doc, {key: _format_value(value) for key, value in payload.items()})
    doc.add_paragraph()


def _render_docx_matrix_section(doc: Document, title: str, headers: list[str], rows: list[list[Any]]) -> None:
    _render_docx_matrix_table(doc, headers, rows, heading=title, heading_level=3)


def _render_main_record_docx(
    doc: Document,
    record: dict,
    *,
    figures: dict | None = None,
    used_figures: set[str] | None = None,
) -> None:
    if used_figures is None:
        used_figures = set()
    doc.add_paragraph(_record_title(record), style="Heading 2")

    matched_figures = select_record_figures(record, figures, used_figures)
    for index, (caption, png_bytes) in enumerate(matched_figures, start=1):
        doc.add_paragraph(normalize_report_text(f"Figure {index}: {caption}"), style="Heading 3")
        try:
            img_stream = io.BytesIO(png_bytes)
            doc.add_picture(img_stream, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            doc.add_paragraph("Figure could not be embedded and was skipped.")
        doc.add_paragraph()

    key_results = _record_key_results(record)
    if key_results:
        _render_record_mapping(doc, "Key Results", key_results)

    for title, payload in _record_main_sections(record):
        _render_record_mapping(doc, title, payload)

    major_events = _tga_major_events(record)
    if major_events:
        _render_docx_matrix_section(
            doc,
            "Major Decomposition Events",
            ["Event", "Midpoint Temperature (°C)", "Mass Loss (%)", "Final Residue (%)"],
            major_events,
        )
    else:
        compact = _record_compact_rows(record)
        if compact:
            headers, rows = compact
            title = "Top Candidates" if str(record.get("analysis_type") or "").upper() == "XRD" else "Compact Key Table"
            _render_docx_matrix_section(doc, title, headers, rows)


def _render_appendix_docx(
    doc: Document,
    *,
    records: list[dict],
    datasets: dict,
    comparison_payload: dict[str, Any] | None,
    figures: dict | None,
) -> None:
    dataset_sections = []
    for dataset_key, dataset in datasets.items():
        payload = _appendix_dataset_metadata(getattr(dataset, "metadata", {}) or {})
        if payload:
            dataset_sections.append((dataset_key, payload))

    record_sections = []
    for record in records:
        sections = _record_appendix_sections(record)
        matrix_sections = _record_appendix_matrix_sections(record)
        dossiers = _xrd_reference_dossiers(record)
        full_rows = _record_full_rows(record)
        if sections or matrix_sections or dossiers or full_rows:
            record_sections.append((record, sections, matrix_sections, dossiers, full_rows))

    comparison_has_content = bool(comparison_payload and (comparison_payload.get("appendix_overview") or comparison_payload.get("appendix_batch_rows")))
    if not (dataset_sections or record_sections or comparison_has_content):
        return

    _add_heading(doc, "Appendix A — Reproducibility and Audit Trail", level=1)

    if dataset_sections:
        doc.add_paragraph("Dataset Import and Metadata Technical Details", style="Heading 2")
        for dataset_key, payload in dataset_sections:
            doc.add_paragraph(_dataset_label(dataset_key, datasets), style="Heading 3")
            _add_key_value_table(doc, payload)
            doc.add_paragraph()

    if comparison_has_content:
        doc.add_paragraph("Comparison Workspace Technical Context", style="Heading 2")
        if comparison_payload.get("appendix_overview"):
            _add_key_value_table(doc, comparison_payload["appendix_overview"])
            doc.add_paragraph()
        batch_rows = comparison_payload.get("appendix_batch_rows") or []
        if batch_rows:
            batch_totals = summarize_batch_outcomes(batch_rows)
            _add_key_value_table(
                doc,
                {
                    "Batch Total": batch_totals["total"],
                    "Saved": batch_totals["saved"],
                    "Blocked": batch_totals["blocked"],
                    "Failed": batch_totals["failed"],
                },
            )
            doc.add_paragraph()
            _render_docx_matrix_section(
                doc,
                "Batch Summary",
                ["Run", "Sample", "Template", "Execution", "Validation", "Calibration", "Reference", "Result ID", "Error ID", "Reason"],
                [
                    [
                        _format_value(row.get("dataset_key")),
                        _format_value(row.get("sample_name")),
                        _format_value(row.get("workflow_template")),
                        _format_value(row.get("execution_status")),
                        _format_value(row.get("validation_status")),
                        _format_value(row.get("calibration_state")),
                        _format_value(row.get("reference_state")),
                        _format_value(row.get("result_id")),
                        _format_value(row.get("error_id")),
                        _format_value(row.get("failure_reason")),
                    ]
                    for row in batch_rows
                ],
            )

    for record, sections, matrix_sections, dossiers, full_rows in record_sections:
        doc.add_paragraph(_record_title(record), style="Heading 2")
        for title, payload in sections:
            _render_record_mapping(doc, title, payload)
        for title, headers, rows in matrix_sections:
            _render_docx_matrix_section(doc, title, headers, rows)
        for dossier in dossiers:
            _render_xrd_reference_dossier_docx(doc, dossier, figures=figures)
        if full_rows:
            headers, rows = full_rows
            _render_docx_matrix_section(doc, "Full Raw Data Table", headers, rows)


def _add_cover_page(doc: Document, branding: dict | None, license_state: dict | None) -> None:
    branding = branding or {}
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = branding.get("report_title") or "MaterialScope Professional Report"
    title_run = title_para.add_run(normalize_report_text(title))
    title_run.bold = True
    title_run.font.size = Pt(20)

    if branding.get("logo_bytes"):
        try:
            doc.add_picture(io.BytesIO(branding["logo_bytes"]), width=Inches(1.4))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_bits = [branding.get("company_name"), branding.get("lab_name")]
    subtitle = " | ".join(bit for bit in subtitle_bits if bit)
    if subtitle:
        subtitle_para.add_run(normalize_report_text(subtitle))

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    analyst = branding.get("analyst_name") or "Analyst not specified"
    meta_para.add_run(
        normalize_report_text(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d  %H:%M')} | Analyst: {analyst}")
    )

    if license_state and license_state.get("status") in {"trial", "activated"}:
        license_para = doc.add_paragraph()
        license_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sku = (license_state.get("license") or {}).get("sku", "Professional")
        license_para.add_run(normalize_report_text(f"License status: {license_state['status']} ({sku})"))

    doc.add_page_break()


def generate_docx_report(
    results: dict,
    datasets: dict,
    figures: Optional[dict] = None,
    file_path_or_buffer: Optional[Union[str, io.BytesIO]] = None,
    branding: Optional[dict] = None,
    comparison_workspace: Optional[dict] = None,
    license_state: Optional[dict] = None,
) -> bytes:
    """Generate a DOCX report from normalized stable/experimental records."""
    valid_results, issues = split_valid_results(results)
    stable_results, experimental_results = partition_results_by_status(valid_results)
    all_records = stable_results + experimental_results
    comparison_payload = _build_comparison_payload(comparison_workspace, datasets, all_records)
    executive_rows = _build_executive_summary_rows(all_records, datasets, comparison_payload)
    executive_intro = _build_executive_summary_intro(all_records, datasets, comparison_payload)
    final_conclusion = _build_final_conclusion_paragraph(all_records, comparison_payload)
    used_figures: set[str] = set()

    doc = Document()
    _add_cover_page(doc, branding, license_state)

    _add_heading(doc, "Executive Summary", level=1)
    if executive_intro:
        doc.add_paragraph(normalize_report_text(executive_intro))
        doc.add_paragraph()
    if executive_rows:
        _add_results_table(
            doc,
            ["Dataset / Set", "Analysis Type", "Key Metrics", "Scientific Interpretation", "Confidence / Limitation"],
            executive_rows,
        )
    else:
        doc.add_paragraph("No analysis results were available for executive summarization.")
    doc.add_paragraph()

    _add_heading(doc, "Experimental Conditions", level=1)
    if not datasets:
        doc.add_paragraph("No dataset metadata available.")
    else:
        for dataset_key, dataset in datasets.items():
            doc.add_paragraph(_dataset_label(dataset_key, datasets), style="Heading 2")
            payload = _main_conditions_payload(dataset_key, dataset)
            if payload:
                _add_key_value_table(doc, payload)
            else:
                doc.add_paragraph("No reader-facing experimental metadata available.")
            doc.add_paragraph()

    if comparison_payload:
        _add_heading(doc, "Comparison Overview", level=1)
        if comparison_payload.get("overview"):
            _add_key_value_table(doc, comparison_payload["overview"])
            doc.add_paragraph()
        if comparison_payload.get("metric_rows"):
            _add_results_table(doc, comparison_payload["metric_headers"], comparison_payload["metric_rows"])
            doc.add_paragraph()
        if comparison_payload.get("excluded_note"):
            doc.add_paragraph("Comparison Coverage Note", style="Heading 2")
            doc.add_paragraph(normalize_report_text(comparison_payload["excluded_note"]))
            doc.add_paragraph()
        if comparison_payload.get("interpretation"):
            doc.add_paragraph("Comparison Interpretation", style="Heading 2")
            doc.add_paragraph(normalize_report_text(comparison_payload["interpretation"]))
            doc.add_paragraph()
        comparison_figures = select_comparison_figures(
            figures,
            used_figures,
            comparison_workspace=comparison_workspace,
        )
        if comparison_figures:
            doc.add_paragraph("Comparison Figures", style="Heading 2")
            for caption, png_bytes in comparison_figures:
                doc.add_paragraph(normalize_report_text(caption), style="Heading 3")
                try:
                    img_stream = io.BytesIO(png_bytes)
                    doc.add_picture(img_stream, width=Inches(5.5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    doc.add_paragraph("Figure could not be embedded and was skipped.")
                doc.add_paragraph()

    _add_heading(doc, "Stable Analyses", level=1)
    if not stable_results:
        doc.add_paragraph("No stable analysis results available.")
    else:
        for record in stable_results:
            _render_main_record_docx(doc, record, figures=figures, used_figures=used_figures)

    if experimental_results:
        _add_heading(doc, "Experimental Analyses", level=1)
        doc.add_paragraph("These results are included for reference but remain outside the stable workflow guarantee.")
        for record in experimental_results:
            _render_main_record_docx(doc, record, figures=figures, used_figures=used_figures)

    report_notes = (branding or {}).get("report_notes")
    if report_notes:
        _add_heading(doc, "Analyst Notes", level=1)
        doc.add_paragraph(str(report_notes))

    if issues:
        _add_heading(doc, "Skipped Records", level=1)
        for issue in issues:
            doc.add_paragraph(issue, style="List Bullet")

    if figures:
        remaining = [
            (caption, png_bytes)
            for caption, png_bytes in figures.items()
            if caption not in used_figures and not _is_comparison_figure_caption(str(caption))
        ]
        if remaining:
            _add_heading(doc, "Additional Figures", level=1)
            for caption, png_bytes in remaining:
                doc.add_paragraph(normalize_report_text(caption), style="Heading 2")
                try:
                    img_stream = io.BytesIO(png_bytes)
                    doc.add_picture(img_stream, width=Inches(5.5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    doc.add_paragraph("Figure could not be embedded and was skipped.")
                doc.add_paragraph()

    if final_conclusion:
        _add_heading(doc, "Final Conclusion", level=1)
        doc.add_paragraph(normalize_report_text(final_conclusion))
        doc.add_paragraph()

    _render_appendix_docx(doc, records=all_records, datasets=datasets, comparison_payload=comparison_payload, figures=figures)

    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()

    if isinstance(file_path_or_buffer, str):
        with open(file_path_or_buffer, "wb") as fh:
            fh.write(docx_bytes)
    elif isinstance(file_path_or_buffer, io.BytesIO):
        file_path_or_buffer.write(docx_bytes)
        file_path_or_buffer.seek(0)

    return docx_bytes


def _comparison_batch_rows(comparison_workspace: dict | None) -> list[dict]:
    comparison_workspace = comparison_workspace or {}
    return normalize_batch_summary_rows(comparison_workspace.get("batch_summary") or [])


def generate_csv_summary(
    results: dict,
    file_path_or_buffer: Optional[Union[str, io.StringIO]] = None,
) -> str:
    """Generate a flat CSV summary from normalized result records."""
    valid_results, _ = split_valid_results(results)
    flat_rows = flatten_result_records(valid_results)

    fieldnames = [
        "result_id",
        "status",
        "analysis_type",
        "dataset_key",
        "section",
        "row_index",
        "field",
        "value",
    ]

    str_buffer = io.StringIO()
    writer = csv.DictWriter(str_buffer, fieldnames=fieldnames)
    writer.writeheader()
    for row in flat_rows:
        writer.writerow(row)
    csv_str = str_buffer.getvalue()

    if isinstance(file_path_or_buffer, str):
        with open(file_path_or_buffer, "w", newline="", encoding="utf-8") as fh:
            fh.write(csv_str)
    elif isinstance(file_path_or_buffer, io.StringIO):
        file_path_or_buffer.write(csv_str)
        file_path_or_buffer.seek(0)

    return csv_str


def pdf_export_available() -> bool:
    """Return whether reportlab is installed for PDF export."""
    try:  # pragma: no cover - availability depends on environment
        import reportlab  # noqa: F401
    except ImportError:
        return False
    return True


def _configure_pdf_font(styles) -> str | None:
    """Register a Unicode-capable PDF font when available."""
    try:  # pragma: no cover - optional dependency path
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except Exception:
        return None

    candidates: list[tuple[str, str]] = []
    try:  # pragma: no cover - optional dependency path
        import matplotlib

        candidates.append(
            (
                "DejaVuSans",
                os.path.join(matplotlib.get_data_path(), "fonts", "ttf", "DejaVuSans.ttf"),
            )
        )
    except Exception:
        pass

    candidates.extend(
        [
            ("NotoSans", "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf"),
            ("DejaVuSans", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
            ("DejaVuSans", os.path.join("C:\\", "Windows", "Fonts", "DejaVuSans.ttf")),
            ("SegoeUI", os.path.join("C:\\", "Windows", "Fonts", "segoeui.ttf")),
            ("Arial", os.path.join("C:\\", "Windows", "Fonts", "arial.ttf")),
            ("Calibri", os.path.join("C:\\", "Windows", "Fonts", "calibri.ttf")),
        ]
    )
    for font_name, font_path in candidates:
        if not os.path.exists(font_path):
            continue
        try:
            if font_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(font_name, font_path))
            for style_name in ("Normal", "Title", "Heading1", "Heading2", "Heading3", "Heading4"):
                if style_name in styles:
                    styles[style_name].fontName = font_name
            return font_name
        except Exception:
            continue
    return None


def _insert_soft_breaks(value: Any, *, chunk: int = 20, max_chars: int | None = None) -> str:
    text = normalize_report_text(_format_value(value))
    if max_chars and len(text) > max_chars:
        text = f"{text[: max_chars - 1]}…"
    if not text:
        return ""
    tokens = text.split(" ")
    wrapped_tokens: list[str] = []
    for token in tokens:
        if len(token) <= chunk:
            wrapped_tokens.append(token)
            continue
        pieces = [token[i : i + chunk] for i in range(0, len(token), chunk)]
        wrapped_tokens.append("\u200b".join(pieces))
    return " ".join(wrapped_tokens)


def _choose_portrait_or_landscape_table_layout(
    headers: list[str],
    rows: list[list[Any]],
    *,
    portrait_width: float,
) -> str:
    if len(headers) >= 8:
        return "landscape"
    max_lengths: list[int] = []
    for column_index, header in enumerate(headers):
        max_len = len(str(header))
        for row in rows[:40]:
            if column_index >= len(row):
                continue
            max_len = max(max_len, len(_format_value(row[column_index])))
        max_lengths.append(max_len)
    estimated = sum(max(52.0, min(150.0, 4.2 * (length + 2))) for length in max_lengths)
    return "landscape" if estimated > portrait_width else "portrait"


def _build_pdf_kv_table(
    payload: dict[str, Any],
    *,
    available_width: float,
    paragraph_style,
    header_style,
    colors,
    font_name: str | None = None,
    compact: bool = False,
):
    from reportlab.platypus import Paragraph, Table, TableStyle

    col_widths = [available_width * 0.34, available_width * 0.66]
    rows = [[Paragraph("Parameter", header_style), Paragraph("Value", header_style)]]
    for key, value in payload.items():
        rows.append(
            [
                Paragraph(_insert_soft_breaks(key, chunk=24), paragraph_style),
                Paragraph(_insert_soft_breaks(value, chunk=20, max_chars=420), paragraph_style),
            ]
        )
    table = Table(rows, colWidths=col_widths, hAlign="LEFT")
    cell_font_size = 8 if compact else 9
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2E4A62")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D1D5DB")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F7FA")]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTSIZE", (0, 1), (-1, -1), cell_font_size),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                *([("FONTNAME", (0, 0), (-1, -1), font_name)] if font_name else []),
            ]
        )
    )
    return table


def _build_pdf_matrix_table(
    headers: list[str],
    rows: list[list[Any]],
    *,
    available_width: float,
    paragraph_style,
    header_style,
    colors,
    font_name: str | None = None,
    compact: bool = False,
):
    from reportlab.platypus import Paragraph, Table, TableStyle

    if not headers:
        return Table([])
    normalized_rows = rows or []
    max_lengths: list[int] = []
    for column_index, header in enumerate(headers):
        max_len = len(str(header))
        for row in normalized_rows[:50]:
            if column_index >= len(row):
                continue
            max_len = max(max_len, len(_format_value(row[column_index])))
        max_lengths.append(max_len)
    weights = [max(1.0, min(7.0, length / 12.0)) for length in max_lengths]
    weight_total = sum(weights) or float(len(weights))
    col_widths = [available_width * weight / weight_total for weight in weights]

    matrix = [[Paragraph(_insert_soft_breaks(header, chunk=24, max_chars=120), header_style) for header in headers]]
    for row in normalized_rows:
        row_cells = []
        for value in row:
            row_cells.append(Paragraph(_insert_soft_breaks(value, chunk=18, max_chars=280), paragraph_style))
        while len(row_cells) < len(headers):
            row_cells.append(Paragraph("", paragraph_style))
        matrix.append(row_cells)

    table = Table(matrix, colWidths=col_widths, hAlign="LEFT", repeatRows=1)
    cell_font_size = 7 if compact else 8
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2E4A62")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D1D5DB")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F7FA")]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTSIZE", (0, 1), (-1, -1), cell_font_size),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                *([("FONTNAME", (0, 0), (-1, -1), font_name)] if font_name else []),
            ]
        )
    )
    return table


def _pdf_render_sections(record: dict) -> list[tuple[str, dict[str, Any]]]:
    """Return PDF-facing sections with redundant legacy scientific block removed."""
    allowed_titles = {
        "Primary Scientific Interpretation",
        "Evidence Supporting This Interpretation",
        "Literature Comparison",
        "Supporting References",
        "Contradictory or Alternative References",
        "Alternative Explanations",
        "Uncertainty and Methodological Limits",
        "Recommended Follow-Up Literature Checks",
        "Recommended Follow-Up Experiments",
    }
    output = []
    for title, payload in _record_main_sections(record):
        if title not in allowed_titles:
            continue
        output.append((title, payload))
    return output


def generate_pdf_report(
    results: dict,
    datasets: dict,
    figures: Optional[dict] = None,
    file_path_or_buffer: Optional[Union[str, io.BytesIO]] = None,
    branding: Optional[dict] = None,
    comparison_workspace: Optional[dict] = None,
    license_state: Optional[dict] = None,
) -> bytes:
    """Generate a scientific-paper-style PDF report with hardened table layout."""
    try:  # pragma: no cover - optional dependency
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.lib.utils import ImageReader
        from reportlab.platypus import (
            BaseDocTemplate,
            Frame,
            Image,
            NextPageTemplate,
            PageBreak,
            PageTemplate,
            Paragraph,
            Spacer,
        )
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise ImportError("PDF export requires reportlab. Install it with: pip install reportlab") from exc

    valid_results, issues = split_valid_results(results)
    stable_results, experimental_results = partition_results_by_status(valid_results)
    all_records = stable_results + experimental_results
    comparison_payload = _build_comparison_payload(comparison_workspace, datasets, all_records)
    abstract_text, abstract_bullets = _build_pdf_abstract_layout(all_records, datasets)
    final_conclusion = _build_final_conclusion_paragraph(all_records, comparison_payload)

    left_right_margin = 19 * mm
    top_bottom_margin = 20 * mm
    portrait_size = A4
    landscape_size = landscape(A4)
    portrait_width = portrait_size[0] - (2 * left_right_margin)
    portrait_height = portrait_size[1] - (2 * top_bottom_margin)
    landscape_width = landscape_size[0] - (2 * left_right_margin)
    landscape_height = landscape_size[1] - (2 * top_bottom_margin)

    buffer = io.BytesIO()
    doc = BaseDocTemplate(
        buffer,
        pagesize=portrait_size,
        leftMargin=left_right_margin,
        rightMargin=left_right_margin,
        topMargin=top_bottom_margin,
        bottomMargin=top_bottom_margin,
        pageCompression=0,
    )
    doc.addPageTemplates(
        [
            PageTemplate(
                id='Portrait',
                pagesize=portrait_size,
                frames=[Frame(left_right_margin, top_bottom_margin, portrait_width, portrait_height, id='portrait_frame')],
            ),
            PageTemplate(
                id='Landscape',
                pagesize=landscape_size,
                frames=[Frame(left_right_margin, top_bottom_margin, landscape_width, landscape_height, id='landscape_frame')],
            ),
        ]
    )

    styles = getSampleStyleSheet()
    pdf_font_name = _configure_pdf_font(styles)
    body_style = ParagraphStyle(
        'PaperBody',
        parent=styles['Normal'],
        fontName=pdf_font_name or styles['Normal'].fontName,
        fontSize=9.5,
        leading=13,
    )
    small_style = ParagraphStyle('PaperSmall', parent=body_style, fontSize=8, leading=10)
    caption_style = ParagraphStyle('FigureCaption', parent=small_style, alignment=1)
    table_header_style = ParagraphStyle('TableHeader', parent=small_style, textColor=colors.white, fontName=pdf_font_name or small_style.fontName)
    heading1 = ParagraphStyle('PaperH1', parent=styles['Heading1'], fontName=pdf_font_name or styles['Heading1'].fontName)
    heading2 = ParagraphStyle('PaperH2', parent=styles['Heading2'], fontName=pdf_font_name or styles['Heading2'].fontName)
    heading3 = ParagraphStyle('PaperH3', parent=styles['Heading3'], fontName=pdf_font_name or styles['Heading3'].fontName)
    title_style = ParagraphStyle('PaperTitle', parent=styles['Title'], fontName=pdf_font_name or styles['Title'].fontName)

    story = []
    current_template = 'Portrait'

    def ensure_template(template_name: str) -> None:
        nonlocal current_template
        if template_name == current_template:
            return
        story.append(NextPageTemplate(template_name))
        story.append(PageBreak())
        current_template = template_name

    def add_heading(text: str, level: int = 1) -> None:
        if level == 1:
            story.append(Paragraph(normalize_report_text(text), heading1))
        elif level == 2:
            story.append(Paragraph(normalize_report_text(text), heading2))
        else:
            story.append(Paragraph(normalize_report_text(text), heading3))

    def add_kv_table(payload: dict[str, Any], *, width: float, compact: bool = False) -> None:
        if not payload:
            return
        story.append(
            _build_pdf_kv_table(
                payload,
                available_width=width,
                paragraph_style=small_style if compact else body_style,
                header_style=table_header_style,
                colors=colors,
                font_name=pdf_font_name,
                compact=compact,
            )
        )

    def add_matrix_table(headers: list[str], rows: list[list[Any]], *, width: float, compact: bool = False) -> None:
        if not headers:
            return
        story.append(
            _build_pdf_matrix_table(
                headers,
                rows,
                available_width=width,
                paragraph_style=small_style if compact else body_style,
                header_style=table_header_style,
                colors=colors,
                font_name=pdf_font_name,
                compact=compact,
            )
        )

    used_figures: set[str] = set()

    def append_record_discussion(record: dict) -> None:
        add_heading(_paper_record_heading(record, datasets), level=2)

        matched_figures = select_record_figures(record, figures, used_figures)
        for index, (caption, png_bytes) in enumerate(matched_figures, start=1):
            try:
                img_reader = ImageReader(io.BytesIO(png_bytes))
                width_px, height_px = img_reader.getSize()
                if not width_px or not height_px:
                    continue
                max_width = portrait_width
                max_height = portrait_height * 0.36
                scale = min(max_width / float(width_px), max_height / float(height_px))
                image = Image(io.BytesIO(png_bytes), width=float(width_px) * scale, height=float(height_px) * scale)
                story.append(image)
                story.append(Paragraph(normalize_report_text(f"Figure {index}. {caption}"), caption_style))
                story.append(Spacer(1, 6))
            except Exception:
                continue

        for title, payload in _pdf_render_sections(record):
            add_heading(title, level=3)
            for key, value in payload.items():
                text = _presentation_text(value)
                if title in {
                    "Evidence Supporting This Interpretation",
                    "Uncertainty and Methodological Limits",
                    "Literature Comparison",
                    "Supporting References",
                    "Contradictory or Alternative References",
                    "Relevant References",
                    "Alternative or Non-Validating References",
                    "Recommended Follow-Up Literature Checks",
                }:
                    text = _presentation_text(f"{key}: {value}")
                story.append(Paragraph(_presentation_text(f"• {text}"), body_style))
            story.append(Spacer(1, 3))

        mini_table = _record_main_mini_table(record)
        if mini_table:
            headers, rows = mini_table
            add_matrix_table(headers, rows, width=portrait_width * 0.6, compact=True)
            story.append(Spacer(1, 4))

    title = (branding or {}).get('report_title') or 'MaterialScope Scientific Report'
    story.append(Paragraph(normalize_report_text(title), title_style))
    story.append(Paragraph(normalize_report_text(datetime.datetime.now().strftime('Generated: %Y-%m-%d %H:%M')), small_style))
    if branding:
        header_bits = [branding.get('company_name'), branding.get('lab_name'), branding.get('analyst_name')]
        meta_line = ' | '.join(bit for bit in header_bits if bit)
        if meta_line:
            story.append(Paragraph(normalize_report_text(meta_line), small_style))
    if license_state and license_state.get('status'):
        story.append(Paragraph(normalize_report_text(f"License: {license_state['status']}"), small_style))
    story.append(Spacer(1, 10))

    if branding and branding.get('logo_bytes'):
        try:
            logo = Image(io.BytesIO(branding['logo_bytes']))
            logo._restrictSize(portrait_width * 0.25, portrait_height * 0.12)
            story.append(logo)
            story.append(Spacer(1, 10))
        except Exception:
            pass

    add_heading('Abstract', level=1)
    story.append(Paragraph(normalize_report_text(abstract_text), body_style))
    if abstract_bullets:
        for bullet in abstract_bullets:
            story.append(Paragraph(normalize_report_text(f"• {bullet}"), body_style))
    story.append(Spacer(1, 10))

    add_heading('Experimental', level=1)
    if not datasets:
        story.append(Paragraph(normalize_report_text('No dataset metadata were available for experimental reporting.'), body_style))
    else:
        for dataset_key, dataset in datasets.items():
            add_heading(_paper_display_label(dataset_key, datasets), level=2)
            story.append(Paragraph(_build_experimental_prose_block(dataset_key, dataset, datasets), body_style))
            story.append(Spacer(1, 4))
    story.append(Spacer(1, 8))

    add_heading('Results and Discussion', level=1)
    if comparison_payload:
        add_heading(f"{comparison_payload.get('overview', {}).get('Analysis Type', 'Cross-Dataset')} Comparison", level=2)
        if comparison_payload.get('overview'):
            compared = comparison_payload.get('overview', {}).get('Compared Datasets') or 'selected datasets'
            figure_ref = comparison_payload.get('overview', {}).get('Saved Figure') or 'Not recorded'
            story.append(
                Paragraph(
                    normalize_report_text(
                        f"Comparison context includes {compared}. Saved comparison figure reference: {figure_ref}."
                    ),
                    body_style,
                )
            )
            story.append(Spacer(1, 3))
        comparison_figures = select_comparison_figures(
            figures,
            used_figures,
            comparison_workspace=comparison_workspace,
        )
        for index, (caption, png_bytes) in enumerate(comparison_figures, start=1):
            try:
                img_reader = ImageReader(io.BytesIO(png_bytes))
                width_px, height_px = img_reader.getSize()
                if not width_px or not height_px:
                    continue
                max_width = portrait_width
                max_height = portrait_height * 0.35
                scale = min(max_width / float(width_px), max_height / float(height_px))
                image = Image(io.BytesIO(png_bytes), width=float(width_px) * scale, height=float(height_px) * scale)
                story.append(image)
                story.append(Paragraph(normalize_report_text(f"Comparison Figure {index}. {caption}"), caption_style))
                story.append(Spacer(1, 4))
            except Exception:
                continue
        if comparison_payload.get('metric_rows'):
            add_matrix_table(comparison_payload['metric_headers'], comparison_payload['metric_rows'], width=portrait_width * 0.85, compact=True)
            story.append(Spacer(1, 4))
        if comparison_payload.get('excluded_note'):
            story.append(Paragraph(normalize_report_text(comparison_payload['excluded_note']), body_style))
        if comparison_payload.get('interpretation'):
            story.append(Paragraph(normalize_report_text(comparison_payload['interpretation']), body_style))
        story.append(Spacer(1, 6))

    if stable_results:
        for record in stable_results:
            append_record_discussion(record)
    else:
        story.append(Paragraph(normalize_report_text('No stable analysis results were available.'), body_style))

    if experimental_results:
        add_heading('Exploratory Analyses (Experimental)', level=2)
        story.append(Paragraph(normalize_report_text('These analyses remain exploratory and should be interpreted with additional caution.'), body_style))
        for record in experimental_results:
            append_record_discussion(record)

    if figures:
        remaining = [
            (caption, png_bytes)
            for caption, png_bytes in figures.items()
            if caption not in used_figures and not _is_comparison_figure_caption(str(caption))
        ]
        if remaining:
            add_heading('Additional Figures', level=2)
        for index, (caption, png_bytes) in enumerate(remaining, start=1):
            try:
                img_reader = ImageReader(io.BytesIO(png_bytes))
                width_px, height_px = img_reader.getSize()
                if not width_px or not height_px:
                    continue
                max_width = portrait_width
                max_height = portrait_height * 0.35
                scale = min(max_width / float(width_px), max_height / float(height_px))
                image = Image(io.BytesIO(png_bytes), width=float(width_px) * scale, height=float(height_px) * scale)
                story.append(image)
                story.append(Paragraph(normalize_report_text(f'Figure {index}. {caption}'), caption_style))
                story.append(Spacer(1, 6))
            except Exception:
                continue

    add_heading('Conclusion', level=1)
    story.append(Paragraph(normalize_report_text(final_conclusion), body_style))

    if issues:
        add_heading('Data and Record Notes', level=2)
        for issue in issues:
            story.append(Paragraph(normalize_report_text(f'• {issue}'), small_style))

    dataset_sections = []
    for dataset_key, dataset in datasets.items():
        payload = _appendix_dataset_metadata(getattr(dataset, 'metadata', {}) or {})
        if payload:
            dataset_sections.append((dataset_key, payload))

    record_sections = []
    for record in all_records:
        sections = _record_appendix_sections(record)
        matrix_sections = _record_appendix_matrix_sections(record)
        dossiers = _xrd_reference_dossiers(record)
        full_rows = _record_full_rows(record)
        if sections or matrix_sections or dossiers or full_rows:
            record_sections.append((record, sections, matrix_sections, dossiers, full_rows))

    comparison_has_content = bool(comparison_payload and (comparison_payload.get('appendix_overview') or comparison_payload.get('appendix_batch_rows')))
    if dataset_sections or record_sections or comparison_has_content:
        story.append(PageBreak())
        ensure_template('Portrait')
        add_heading('Supplementary Technical Record (Appendix A)', level=1)

        if dataset_sections:
            add_heading('Dataset Import and Metadata Technical Details', level=2)
            for dataset_key, payload in dataset_sections:
                ensure_template('Portrait')
                add_heading(_dataset_label(dataset_key, datasets), level=3)
                add_kv_table(payload, width=portrait_width, compact=True)
                story.append(Spacer(1, 4))

        if comparison_has_content:
            add_heading('Comparison Workspace Technical Context', level=2)
            if comparison_payload.get('appendix_overview'):
                ensure_template('Portrait')
                add_kv_table(comparison_payload['appendix_overview'], width=portrait_width, compact=True)
                story.append(Spacer(1, 4))
            batch_rows = comparison_payload.get('appendix_batch_rows') or []
            if batch_rows:
                batch_totals = summarize_batch_outcomes(batch_rows)
                ensure_template('Portrait')
                add_kv_table(
                    {
                        'Batch Total': batch_totals['total'],
                        'Saved': batch_totals['saved'],
                        'Blocked': batch_totals['blocked'],
                        'Failed': batch_totals['failed'],
                    },
                    width=portrait_width,
                    compact=True,
                )
                story.append(Spacer(1, 4))
                batch_headers = ['Run', 'Sample', 'Template', 'Execution', 'Validation', 'Calibration', 'Reference', 'Result ID', 'Error ID', 'Reason']
                batch_matrix = [
                    [
                        _format_value(row.get('dataset_key')),
                        _format_value(row.get('sample_name')),
                        _format_value(row.get('workflow_template')),
                        _format_value(row.get('execution_status')),
                        _format_value(row.get('validation_status')),
                        _format_value(row.get('calibration_state')),
                        _format_value(row.get('reference_state')),
                        _format_value(row.get('result_id')),
                        _format_value(row.get('error_id')),
                        _format_value(row.get('failure_reason')),
                    ]
                    for row in batch_rows
                ]
                batch_layout = _choose_portrait_or_landscape_table_layout(batch_headers, batch_matrix, portrait_width=portrait_width)
                ensure_template('Landscape' if batch_layout == 'landscape' else 'Portrait')
                add_matrix_table(batch_headers, batch_matrix, width=landscape_width if batch_layout == 'landscape' else portrait_width, compact=True)
                story.append(Spacer(1, 4))

        for record, sections, matrix_sections, dossiers, full_rows in record_sections:
            ensure_template('Portrait')
            add_heading(_record_title(record), level=2)
            for title, payload in sections:
                add_heading(title, level=3)
                add_kv_table(payload, width=portrait_width, compact=True)
                story.append(Spacer(1, 4))
            for title, headers, rows in matrix_sections:
                matrix_layout = _choose_portrait_or_landscape_table_layout(headers, rows, portrait_width=portrait_width)
                ensure_template('Landscape' if matrix_layout == 'landscape' else 'Portrait')
                add_heading(title, level=3)
                add_matrix_table(headers, rows, width=landscape_width if matrix_layout == 'landscape' else portrait_width, compact=True)
                story.append(Spacer(1, 4))
            for dossier in dossiers:
                _render_xrd_reference_dossier_pdf(
                    story,
                    dossier,
                    figures=figures,
                    add_heading=add_heading,
                    add_kv_table=add_kv_table,
                    add_matrix_table=add_matrix_table,
                    ensure_template=ensure_template,
                    portrait_width=portrait_width,
                    landscape_width=landscape_width,
                    body_style=body_style,
                    small_style=small_style,
                    caption_style=caption_style,
                    paragraph_factory=Paragraph,
                    image_factory=Image,
                    spacer_factory=Spacer,
                )
            if full_rows:
                headers, rows = full_rows
                raw_layout = _choose_portrait_or_landscape_table_layout(headers, rows, portrait_width=portrait_width)
                ensure_template('Landscape' if raw_layout == 'landscape' else 'Portrait')
                add_heading('Full Raw Data Table', level=3)
                add_matrix_table(headers, rows, width=landscape_width if raw_layout == 'landscape' else portrait_width, compact=True)
                story.append(Spacer(1, 4))

    doc.build(story)
    pdf_bytes = buffer.getvalue()

    if isinstance(file_path_or_buffer, str):
        with open(file_path_or_buffer, 'wb') as fh:
            fh.write(pdf_bytes)
    elif isinstance(file_path_or_buffer, io.BytesIO):
        file_path_or_buffer.write(pdf_bytes)
        file_path_or_buffer.seek(0)

    return pdf_bytes
